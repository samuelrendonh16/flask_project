from flask import Flask, jsonify, request, render_template, send_file, url_for, redirect, session, current_app, make_response, abort
from flask_cors import CORS, cross_origin
from flask_mail import Mail, Message
from db import db, init_db, Bodegas, Licencia, Usuarios, Consecutivos, Entradas1, Entradas2, Referencia, Salidas1, Salidas2, Grupo, Traslados1, Traslados2, SaldosBodega, Unidades, SubGrupos, Subcategorias, EstadoProducto, Proveedores, OrdenesCompra1, OrdenesCompra2, Paises, Departamentos, Ciudades, Bancos, Compras1, Compras2, CentroCostos, Permisos, UsuariosPermisos
from datetime import datetime, timedelta, date
import traceback
import hashlib
import logging
import random
import string
import base64
from PIL import Image, ImageDraw, ImageFont
import io
from sqlalchemy.exc import SQLAlchemyError
from sqlalchemy import text
from sqlalchemy.exc import IntegrityError
from decimal import Decimal, InvalidOperation
from sqlalchemy import func, desc
from flask_sqlalchemy import SQLAlchemy
import uuid
from sqlalchemy.dialects.postgresql import UUID
from dateutil import parser
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill
from openpyxl.utils import get_column_letter
from sqlalchemy.orm.exc import NoResultFound
from flask import after_this_request
import tempfile
from functools import wraps
import os
from flask_login import login_required, login_user, current_user
from werkzeug.security import generate_password_hash, check_password_hash
from sqlalchemy.orm import joinedload
from docx import Document
from docx.shared import Inches
from werkzeug.exceptions import InternalServerError
from io import BytesIO
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, Cm
from flask import jsonify, request, redirect, url_for, session
from flask_login import logout_user
from werkzeug.security import generate_password_hash
from sqlalchemy.exc import SQLAlchemyError
from functools import wraps
import logging
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
import time
from flask_login import LoginManager
from sqlalchemy.exc import ProgrammingError, OperationalError
from sqlalchemy import inspect, text

def generar_password():
    return ''.join(random.choices(string.ascii_letters + string.digits, k=10))

logging.basicConfig(level=logging.DEBUG, format='%(asctime)s - %(levelname)s - %(message)s')
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

app = Flask(__name__)
app.secret_key = os.environ.get('FLASK_SECRET_KEY') or '8f42a73d4f8c4c1d9c525f68995ced6b2537a quaracer6492'
CORS(app, resources={r"/api/*": {
    "origins": ["https://migsistemasweb.com", "https://www.migsistemasweb.com"],
    "methods": ["GET", "POST", "PUT", "DELETE", "OPTIONS"],
    "allow_headers": ["Content-Type", "Authorization"],
    "supports_credentials": True
}})
# Configuración de la base de datos
app.config['SQLALCHEMY_DATABASE_URI'] = 'postgresql://postgres:1234@194.163.45.32/INVENTARIO'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['SESSION_COOKIE_DOMAIN'] = '.migsistemasweb.com'

init_db(app)

# Configuración del logging
logging.basicConfig(level=logging.DEBUG, format='%(asctime)s - %(levelname)s - %(message)s')

# Configuración del correo
app.config['MAIL_SERVER'] = 'smtp.gmail.com'
app.config['MAIL_PORT'] = 587
app.config['MAIL_USE_TLS'] = True
app.config['MAIL_USE_SSL'] = False
app.config['MAIL_USERNAME'] = 'info@migsistemas.com'
app.config['MAIL_PASSWORD'] = 'eejgelzxmgrfrkdh'  # Esta es la contraseña de aplicación
app.config['MAIL_DEFAULT_SENDER'] = 'info@migsistemas.com'
mail = Mail(app)

UPLOAD_FOLDER = '/var/www/html/flask_project/TEMP'
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

# Verificar la conexión a la base de datos
try:
    with app.app_context():
        db.session.execute(text('SELECT 1'))
    logging.info("Conexión a la base de datos exitosa")
except SQLAlchemyError as e:
    logging.error(f"Error al conectar a la base de datos: {str(e)}")

# Crear las tablas si no existen
with app.app_context():
    db.create_all()

@app.route('/api/inventario')
def get_inventario_data():
    inventario_items = [
        # {"icon": "icon-ordenes", "text": "Órdenes de compras"},
        # {"icon": "icon-compras", "text": "Compras proveedor"},
        # {"icon": "icon-traslados", "text": "Traslados a Bodegas"},
        {"icon": "icon-entradas", "text": "Entradas a Bodegas"},
        {"icon": "icon-salidas", "text": "Salidas a Bodegas"},
        {"icon": "icon-consulta", "text": "Consulta Inventario"}
        # {"icon": "icon-maxmin", "text": "Máximos y Mínimos"},
        # {"icon": "icon-fisico", "text": "Inventario Físico"}
    ]
    return jsonify(inventario_items)

@app.route('/api/referencias')
def get_referencias():
    filtro = request.args.get('filtro', '')
    referencias = Referencia.query.filter(
        (Referencia.IdReferencia.ilike(f'%{filtro}%')) | 
        (Referencia.Referencia.ilike(f'%{filtro}%'))
    ).all()
    return jsonify([{
        'IdReferencia': ref.IdReferencia,
        'Referencia': ref.Referencia,
        'PrecioVenta1': str(ref.PrecioVenta1),
        'IVA': str(ref.IVA),
        'Ubicacion': ref.Ubicacion,
        'idbodega': ref.idbodega,
        'IdUnidad': ref.IdUnidad,
    } for ref in referencias])

@app.route('/api/buscar_productos_editar', methods=['GET'])
def buscar_productos_editar():
    busqueda = request.args.get('buscar', '')
    app.logger.info(f"Búsqueda recibida: {busqueda}")

    try:
        referencias = Referencia.query.filter(
            (Referencia.IdReferencia.ilike(f'%{busqueda}%')) | 
            (Referencia.Referencia.ilike(f'%{busqueda}%'))
        ).limit(50).all()
        
        app.logger.info(f"Resultados encontrados: {len(referencias)}")
        
        resultado = [{
            'IdReferencia': ref.IdReferencia,
            'Referencia': ref.Referencia,
            'IdGrupo': ref.IdGrupo
        } for ref in referencias]
        
        return jsonify(resultado)
    except Exception as e:
        app.logger.error(f"Error en la búsqueda: {str(e)}")
        return jsonify({'error': 'Error en la búsqueda de productos'}), 500

# Ruta de prueba para verificar que el servidor está respondiendo
@app.route('/api/test', methods=['GET'])
def test_api():
    return jsonify({"message": "API is working"}), 200

@app.route('/api/buscar_productos_editar/<string:id>', methods=['GET'])
def obtener_producto_editar(id):
    try:
        referencia = Referencia.query.get_or_404(id)
        return jsonify({
            'IdReferencia': referencia.IdReferencia,
            'Referencia': referencia.Referencia,
            'IdGrupo': referencia.IdGrupo,
            'idsubgrupo': referencia.idsubgrupo,
            'idsubcategoria': referencia.idsubcategoria,
            'IdUnidad': referencia.IdUnidad,
            'idbodega': referencia.idbodega,
            'Costo': str(referencia.Costo),
            'PrecioVenta1': str(referencia.PrecioVenta1),
            'IVA': str(referencia.IVA),
            'Ubicacion': referencia.Ubicacion,
            'Marca': referencia.Marca,
            'EstadoProducto': referencia.EstadoProducto,
            'Estado': referencia.Estado,
            'Tipo': referencia.Tipo,
            'ManejaInventario': referencia.ManejaInventario,
            'productoagotado': referencia.productoagotado,
            'modificaprecio': referencia.modificaprecio,
            'compuesto': getattr(referencia, 'compuesto', False)  # Asumiendo que este campo existe
        })
    except Exception as e:
        app.logger.error(f"Error al obtener producto: {str(e)}")
        return jsonify({'error': 'Error al obtener detalles del producto'}), 500

@app.route('/api/maestros')
def get_maestros_data():
    maestros_items = [
        {
            "icon": "icon-terceros",
            "text": "Terceros",
            "subitems": [
                {"icon": "icon-clientes", "text": "Clientes"},
                {"icon": "icon-proveedores", "text": "Proveedores"},
                {"icon": "icon-vendedores", "text": "Vendedores - Meseros - Empleados"}
            ]
        },
        {
            "icon": "icon-productos",
            "text": "Productos",
            "subitems": [
                {"icon": "icon-grupos", "text": "Grupos - Familias - Categorías"},
                {"icon": "icon-articulos", "text": "Productos - Artículos - Referencias"},
                {"icon": "icon-subgrupos", "text": "SubGrupos"},
                {"icon": "icon-subcategorias", "text": "SubCategorías"},
                {"icon": "icon-lineas", "text": "Líneas"},
                {"icon": "icon-comentarios", "text": "Grupos comentarios"},
                {"icon": "icon-descuentos", "text": "Descuentos"},
                {"icon": "icon-unidades", "text": "Unidad de medidas"},
                {"icon": "icon-conectores", "text": "Conectores"}
            ]
        },
        {
            "icon": "icon-otros",
            "text": "Otros",
            "subitems": [
                {"icon": "icon-bodegas", "text": "Bodegas"}
            ]
        },
    ]
    return jsonify(maestros_items)

@app.route('/api/consulta_inventario')
def get_consulta_inventario():
    # Obtiene la fecha actual y el mes actual en el formato 'YYYYMM'
    mes_actual = datetime.now().strftime('%Y%m')

    referencias = db.session.query(
        Referencia, 
        func.sum(func.coalesce(SaldosBodega.Saldo, 0)).label('Saldo')
    ).outerjoin(
        SaldosBodega, 
        (Referencia.IdReferencia == SaldosBodega.IdReferencia)
    ).filter(
        Referencia.Estado == True
    ).group_by(
        Referencia.IdReferencia
    ).all()

    return jsonify([{
        'IDReferencia': ref.IdReferencia,
        'Referencia': ref.Referencia,
        'Marca': ref.Marca or '',
        'Precio_Venta': str(ref.PrecioVenta1) if ref.PrecioVenta1 else '',
        'Ubicación': ref.Ubicacion or '',
        'Grupo': ref.IdGrupo or '',
        'ID_Unidad': ref.IdUnidad or '',
        'Bodega': ref.idbodega or '',
        'Saldo': str(saldo),  # Aquí se incluye el saldo unificado
        'Costo': str(ref.Costo) if ref.Costo else '',
        'EstadoProducto': ref.EstadoProducto or ''
    } for ref, saldo in referencias])

@app.route('/api/bodegas', methods=['GET', 'POST', 'PUT'])
def manejar_bodegas():
    if request.method == 'POST':
        data = request.json
        nueva_bodega = Bodegas(
            IdBodega=data['IdBodega'],
            Descripcion=data['Descripcion'],
            Estado=data['Estado'],
            Email=data.get('Email'),
            nombrepunto=data.get('nombrepunto'),
            direccionpunto=data.get('direccionpunto'),
            telefonopunto=data.get('telefonopunto'),
            Encargado=data.get('Encargado')  # Corregido a 'Encargado'
        )
        db.session.add(nueva_bodega)
        db.session.commit()
        return jsonify({'message': 'Bodega creada exitosamente'}), 201
    elif request.method == 'PUT':
        data = request.json
        bodega = Bodegas.query.get(data['IdBodega'])
        if bodega:
            bodega.Descripcion = data['Descripcion']
            bodega.Estado = data['Estado']
            bodega.Email = data.get('Email')
            bodega.nombrepunto = data.get('nombrepunto')
            bodega.direccionpunto = data.get('direccionpunto')
            bodega.telefonopunto = data.get('telefonopunto')
            bodega.Encargado = data.get('Encargado')  # Corregido a 'Encargado'
            db.session.commit()
            return jsonify({'message': 'Bodega actualizada exitosamente'}), 200
        else:
            return jsonify({'message': 'Bodega no encontrada'}), 404
    else:
        bodegas = Bodegas.query.all()
        return jsonify([{
            'IdBodega': bodega.IdBodega,
            'Descripcion': bodega.Descripcion,
            'Estado': bodega.Estado,
            'Email': bodega.Email,
            'nombrepunto': bodega.nombrepunto,
            'direccionpunto': bodega.direccionpunto,
            'telefonopunto': bodega.telefonopunto,
            'Encargado': bodega.Encargado  # Corregido a 'Encargado'
        } for bodega in bodegas])


@app.route('/api/proveedores', methods=['POST', 'PUT'])
def manejar_proveedores():
    data = request.json
    app.logger.info(f"Datos recibidos: {data}")
   
    proveedor_data = {
        'Nit': data.get('Nit'),
        'DV': data.get('DV'),
        'RazonSocial': data.get('RazonSocial'),
        'Nombre1': data.get('Nombre1'),
        'Nombre2': data.get('Nombre2'),
        'Apellido1': data.get('Apellido1'),
        'Apellido2': data.get('Apellido2'),
        'Email': data.get('Email'),
        'Cuenta': data.get('Cuenta'),
        'CxP': data.get('CxP'),
        'DiasCredito': data.get('DiasCredito'),
        'Estado': data.get('Estado', True),
        'IdDepartamento': data.get('IdDepartamento'),
        'IdCiudad': data.get('IdCiudad'),
        'fechacreacion': datetime.now()
    }
    try:
        proveedor = Proveedores.query.get(proveedor_data['Nit'])
        if proveedor and request.method == 'POST':
            app.logger.warning(f"Intento de crear un proveedor existente: {proveedor_data['Nit']}")
            return jsonify({'error': 'El proveedor ya existe'}), 400
        elif proveedor and request.method == 'PUT':
            app.logger.info(f"Actualizando proveedor existente: {proveedor_data['Nit']}")
            for key, value in proveedor_data.items():
                if value is not None:  # Solo actualiza si el valor no es None
                    setattr(proveedor, key, value)
        else:
            app.logger.info(f"Creando nuevo proveedor: {proveedor_data['Nit']}")
            proveedor = Proveedores(**proveedor_data)
            db.session.add(proveedor)
        db.session.commit()
        app.logger.info(f"Proveedor guardado exitosamente: {proveedor_data['Nit']}")
        return jsonify({'message': 'Proveedor guardado exitosamente'}), 200
    except SQLAlchemyError as e:
        db.session.rollback()
        app.logger.error(f"Error de SQLAlchemy al guardar proveedor: {str(e)}")
        return jsonify({'error': f'Error de base de datos: {str(e)}'}), 500
    except Exception as e:
        db.session.rollback()
        app.logger.error(f"Error inesperado al guardar proveedor: {str(e)}")
        return jsonify({'error': f'Error inesperado: {str(e)}'}), 500

@app.route('/api/proveedores', methods=['GET'])
def obtener_proveedores():
    try:
        proveedores = Proveedores.query.all()
        return jsonify([{
            'Nit': p.Nit,
            'DV': p.DV,
            'RazonSocial': p.RazonSocial,
            'Nombre1': p.Nombre1,
            'Nombre2': p.Nombre2,
            'Apellido1': p.Apellido1,
            'Apellido2': p.Apellido2,
            'Email': p.Email,
            'Cuenta': p.Cuenta,
            'CxP': p.CxP,
            'DiasCredito': p.DiasCredito,
            'Estado': p.Estado,
            'fechacreacion': p.fechacreacion.isoformat() if p.fechacreacion else None
        } for p in proveedores])
    except Exception as e:
        app.logger.error(f"Error al obtener proveedores: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/proveedores/<string:nit>', methods=['DELETE'])
def eliminar_proveedor(nit):
    try:
        proveedor = Proveedores.query.get(nit)
        if proveedor:
            db.session.delete(proveedor)
            db.session.commit()
            return jsonify({'message': 'Proveedor eliminado exitosamente'}), 200
        else:
            return jsonify({'error': 'Proveedor no encontrado'}), 404
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500
        
@app.route('/api/permisos', methods=['GET', 'POST'])
def manejar_permisos():
    if request.method == 'GET':
        permisos = Permisos.query.all()
        return jsonify([{"IdPermiso": p.IdPermiso, "NombrePermiso": p.NombrePermiso, "Descripcion": p.Descripcion} for p in permisos])
    elif request.method == 'POST':
        data = request.json
        nuevo_permiso = Permisos(NombrePermiso=data['NombrePermiso'], Descripcion=data.get('Descripcion'))
        db.session.add(nuevo_permiso)
        db.session.commit()
        return jsonify({"message": "Permiso creado exitosamente"}), 201

@app.route('/api/usuarios/<string:id_usuario>/permisos', methods=['GET', 'POST'])
def manejar_permisos_usuario(id_usuario):
    if request.method == 'GET':
        permisos = db.session.query(Permisos).join(UsuariosPermisos).filter(UsuariosPermisos.IdUsuario == id_usuario).all()
        return jsonify([{"IdPermiso": p.IdPermiso, "NombrePermiso": p.NombrePermiso} for p in permisos])
    elif request.method == 'POST':
        data = request.json
        nuevo_permiso_usuario = UsuariosPermisos(IdUsuario=id_usuario, IdPermiso=data['IdPermiso'])
        db.session.add(nuevo_permiso_usuario)
        db.session.commit()
        return jsonify({"message": "Permiso asignado exitosamente"}), 201

@app.route('/api/estado_producto', methods=['GET', 'POST'])
def manejar_estado_producto():
    if request.method == 'GET':
        try:
            estados = EstadoProducto.query.all()
            return jsonify([{
                'IdEstadoProducto': estado.IdEstadoProducto,
                'EstadoProducto': estado.EstadoProducto,
                'Estado': estado.Estado
            } for estado in estados])
        except Exception as e:
            return jsonify({'success': False, 'message': str(e)}), 500

    elif request.method == 'POST':
        try:
            data = request.json
            nuevo_estado = EstadoProducto(
                IdEstadoProducto=data['IdEstadoProducto'],
                EstadoProducto=data['EstadoProducto'],
                Estado=data['Estado']
            )
            db.session.add(nuevo_estado)
            db.session.commit()
            return jsonify({'success': True, 'message': 'Estado de Producto creado exitosamente'})
        except Exception as e:
            db.session.rollback()
            return jsonify({'success': False, 'message': str(e)}), 400

@app.route('/api/estado_producto/<string:id>', methods=['PUT', 'DELETE'])
def manejar_estado_producto_individual(id):
    estado = EstadoProducto.query.get(id)
    if not estado:
        return jsonify({'success': False, 'message': 'Estado de Producto no encontrado'}), 404

    if request.method == 'PUT':
        try:
            data = request.json
            estado.EstadoProducto = data['EstadoProducto']
            estado.Estado = data['Estado']
            db.session.commit()
            return jsonify({'success': True, 'message': 'Estado de Producto actualizado exitosamente'})
        except Exception as e:
            db.session.rollback()
            return jsonify({'success': False, 'message': str(e)}), 400

    elif request.method == 'DELETE':
        try:
            db.session.delete(estado)
            db.session.commit()
            return jsonify({'success': True, 'message': 'Estado de Producto eliminado exitosamente'})
        except Exception as e:
            db.session.rollback()
            return jsonify({'success': False, 'message': str(e)}), 400

@app.route('/api/estado_producto', methods=['POST'])
def crear_estado_producto():
    data = request.json
    try:
        nuevo_estado = EstadoProducto(
            IdEstadoProducto=data['IdEstadoProducto'],
            EstadoProducto=data['EstadoProducto'],
            Estado=data['Estado']
        )
        db.session.add(nuevo_estado)
        db.session.commit()
        return jsonify({'success': True, 'message': 'Estado de Producto creado exitosamente'}), 201
    except IntegrityError as e:
        db.session.rollback()
        return jsonify({'success': False, 'message': 'Error de integridad: ' + str(e)}), 400
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'message': 'Error: ' + str(e)}), 500

@app.route('/api/estado_producto/<string:id>', methods=['DELETE'])
def eliminar_estado_producto(id):
    estado = EstadoProducto.query.get(id)
    if estado:
        db.session.delete(estado)
        try:
            db.session.commit()
            return jsonify({'success': True, 'message': 'Estado de Producto eliminado exitosamente'})
        except Exception as e:
            db.session.rollback()
            return jsonify({'success': False, 'message': str(e)}), 400
    else:
        return jsonify({'success': False, 'message': 'Estado de Producto no encontrado'}), 404

@app.route('/api/unidades', methods=['POST'])
def crear_unidad():
    data = request.json
    
    if not data or 'IdUnidad' not in data or 'Unidad' not in data:
        return jsonify({'error': 'Datos incompletos'}), 400

    nueva_unidad = Unidades(
        IdUnidad=data['IdUnidad'],
        Unidad=data['Unidad'],
        Estado=data.get('Estado', True)
    )

    try:
        db.session.add(nueva_unidad)
        db.session.commit()
        return jsonify({
            'message': 'Unidad creada exitosamente',
            'unidad': {
                'IdUnidad': nueva_unidad.IdUnidad,
                'Unidad': nueva_unidad.Unidad,
                'Estado': nueva_unidad.Estado
            }
        }), 201
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500

@app.route('/api/unidades', methods=['GET'])
def obtener_unidades():
    unidades = Unidades.query.all()
    return jsonify([{
        'IdUnidad': u.IdUnidad,
        'Unidad': u.Unidad,
        'Estado': u.Estado
    } for u in unidades])

@app.route('/api/unidades/activas', methods=['GET'])
def obtener_unidades_activas():
    unidades = Unidades.query.filter_by(Estado=True).all()
    return jsonify([{
        'IdUnidad': u.IdUnidad,
        'Unidad': u.Unidad
    } for u in unidades])

from flask import request, send_file
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill
from openpyxl.utils import get_column_letter
import io
from datetime import datetime

@app.route('/api/exportar_salidas_excel')
@cross_origin(supports_credentials=True)
def exportar_salidas_excel():
    try:
        fecha_inicio = request.args.get('fecha_inicio')
        fecha_fin = request.args.get('fecha_fin')
        
        current_app.logger.info(f"Iniciando exportación de salidas: {fecha_inicio} - {fecha_fin}")
        
        if not fecha_inicio or not fecha_fin:
            raise ValueError("Fechas de inicio y fin son requeridas")

        salidas = obtener_datos_salidas(fecha_inicio, fecha_fin)
        
        current_app.logger.info(f"Datos obtenidos: {len(salidas)} salidas")
        
        if not salidas:
            return jsonify({"message": "No hay datos para exportar en el rango de fechas especificado"}), 404

        wb = Workbook()
        ws = wb.active
        ws.title = "Salidas de Inventario"

        # Estilos
        title_font = Font(name='Arial', size=16, bold=True)
        header_font = Font(name='Arial', size=12, bold=True)
        
        # Título
        ws.merge_cells('A1:H1')
        ws['A1'] = "Salidas de Inventario"
        ws['A1'].font = title_font
        ws['A1'].alignment = Alignment(horizontal='center')

        # Información de la empresa
        ws['A3'] = "CCD INGENIERIA Y CONSTRUCCIONES S.A.S."
        ws['A4'] = "901092189-5"
        ws['A5'] = "CCD INGENIERIA Y CONSTRUCCIONES"
        
        ws['F3'] = f"Fecha: {fecha_inicio} - {fecha_fin}"
        ws['F4'] = "Dirección: CR 78 A 45 G G 14"
        ws['F5'] = "Teléfono: 3043821361"

        # Encabezados
        headers = ['Numero', 'Fecha y Hora', 'Codigo', 'Producto', 'Cant.', 'Valor', 'Total', 'Observaciones']
        for col, header in enumerate(headers, start=1):
            cell = ws.cell(row=7, column=col, value=header)
            cell.font = header_font
            cell.fill = PatternFill(start_color="4F81BD", end_color="4F81BD", fill_type="solid")
            cell.alignment = Alignment(horizontal='center')

        # Datos
        for row, salida in enumerate(salidas, start=8):
            ws.cell(row=row, column=1, value=salida['Numero'])
            ws.cell(row=row, column=2, value=salida['Fecha'])
            ws.cell(row=row, column=3, value=salida['IdReferencia'])
            ws.cell(row=row, column=4, value=salida['Descripcion'])
            ws.cell(row=row, column=5, value=salida['Cantidad'])
            ws.cell(row=row, column=6, value=salida['Valor'])
            ws.cell(row=row, column=7, value=salida['Total'])
            ws.cell(row=row, column=8, value=salida.get('Observaciones', ''))

        # Ajustar el ancho de las columnas
        for column in ws.columns:
            max_length = 0
            column_letter = get_column_letter(column[0].column)
            for cell in column:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            adjusted_width = (max_length + 2)
            ws.column_dimensions[column_letter].width = adjusted_width

        current_app.logger.info("Archivo Excel creado exitosamente")
        
        output = io.BytesIO()
        wb.save(output)
        output.seek(0)

        response = make_response(output.getvalue())
        response.headers['Content-Disposition'] = 'attachment; filename=Salidas_de_Inventario.xlsx'
        response.headers['Content-type'] = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        response.headers['Access-Control-Allow-Origin'] = '*'
        
        current_app.logger.info("Archivo Excel enviado exitosamente")
        return response

    except ValueError as ve:
        current_app.logger.error(f"Error de validación: {str(ve)}")
        return jsonify({"error": str(ve)}), 400
    except Exception as e:
        current_app.logger.error(f"Error al exportar salidas: {str(e)}")
        current_app.logger.error(traceback.format_exc())
        return jsonify({"error": "Error interno del servidor", "details": str(e), "traceback": traceback.format_exc()}), 500

@app.route('/api/debug_exportar_salidas_excel')
@cross_origin(supports_credentials=True)
def debug_exportar_salidas_excel():
    try:
        fecha_inicio = request.args.get('fecha_inicio')
        fecha_fin = request.args.get('fecha_fin')
        
        current_app.logger.info(f"Iniciando depuración de exportación de salidas: {fecha_inicio} - {fecha_fin}")
        
        # Paso 1: Obtener datos
        salidas = obtener_datos_salidas(fecha_inicio, fecha_fin)
        current_app.logger.info(f"Datos obtenidos: {len(salidas)} salidas")
        
        if not salidas:
            return jsonify({"message": "No hay datos para exportar en el rango de fechas especificado"}), 404

        # Paso 2: Crear workbook
        wb = Workbook()
        ws = wb.active
        ws.title = "Salidas de Inventario"
        current_app.logger.info("Workbook creado")
        
        # Paso 3: Añadir encabezados
        headers = ['Numero', 'Fecha', 'Codigo', 'Producto', 'Cant.', 'Valor', 'Total']
        for col, header in enumerate(headers, start=1):
            ws.cell(row=1, column=col, value=header)
        current_app.logger.info("Encabezados añadidos")
        
        # Paso 4: Añadir datos
        for row, salida in enumerate(salidas, start=2):
            ws.cell(row=row, column=1, value=salida['Numero'])
            ws.cell(row=row, column=2, value=salida['Fecha'])
            ws.cell(row=row, column=3, value=salida['IdReferencia'])
            ws.cell(row=row, column=4, value=salida['Descripcion'])
            ws.cell(row=row, column=5, value=salida['Cantidad'])
            ws.cell(row=row, column=6, value=salida['Valor'])
            ws.cell(row=row, column=7, value=salida['Total'])
        current_app.logger.info("Datos añadidos")
        
        # Paso 5: Guardar en BytesIO
        output = io.BytesIO()
        wb.save(output)
        output.seek(0)
        current_app.logger.info("Archivo guardado en BytesIO")
        
        # Paso 6: Enviar archivo
        response = make_response(output.getvalue())
        response.headers['Content-Disposition'] = 'attachment; filename=Debug_Salidas_de_Inventario.xlsx'
        response.headers['Content-type'] = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        response.headers['Access-Control-Allow-Origin'] = '*'
        
        current_app.logger.info("Archivo Excel de depuración enviado exitosamente")
        return response

    except ValueError as ve:
        current_app.logger.error(f"Error de validación en depuración: {str(ve)}")
        return jsonify({"error": str(ve)}), 400
    except Exception as e:
        current_app.logger.error(f"Error en depuración de exportación de salidas: {str(e)}")
        current_app.logger.error(traceback.format_exc())
        return jsonify({"error": str(e), "traceback": traceback.format_exc()}), 500

@app.route('/api/paises', methods=['GET'])
def obtener_paises():
    paises = Paises.query.all()
    return jsonify([
        {
            'IdPais': pais.IdPais,
            'Pais': pais.Pais,
            'Estado': pais.Estado
        } for pais in paises
    ])

@app.route('/api/paises', methods=['POST'])
def crear_pais():
    data = request.json
    nuevo_pais = Paises(
        IdPais=data['IdPais'],
        Pais=data['Pais'],
        Estado=data['Estado']
    )
    db.session.add(nuevo_pais)
    db.session.commit()
    return jsonify({'message': 'País creado exitosamente'}), 201

@app.route('/api/paises/<string:id>', methods=['PUT'])
def actualizar_pais(id):
    pais = Paises.query.get(id)
    if not pais:
        return jsonify({'message': 'País no encontrado'}), 404
    data = request.json
    pais.Pais = data['Pais']
    pais.Estado = data['Estado']
    db.session.commit()
    return jsonify({'message': 'País actualizado exitosamente'})

@app.route('/api/exportar_entradas_excel')
@cross_origin(supports_credentials=True)
def exportar_entradas_excel():
    try:
        fecha_inicio = request.args.get('fecha_inicio')
        fecha_fin = request.args.get('fecha_fin')
        
        current_app.logger.info(f"Iniciando exportación de entradas: {fecha_inicio} - {fecha_fin}")
        
        if not fecha_inicio or not fecha_fin:
            raise ValueError("Fechas de inicio y fin son requeridas")

        entradas = obtener_datos_entradas(fecha_inicio, fecha_fin)
        
        current_app.logger.info(f"Datos obtenidos: {len(entradas)} entradas")
        
        if not entradas:
            return jsonify({"message": "No hay datos para exportar en el rango de fechas especificado"}), 404

        wb = Workbook()
        ws = wb.active
        ws.title = "Entradas de Inventario"

        # Estilos
        title_font = Font(name='Arial', size=16, bold=True)
        header_font = Font(name='Arial', size=12, bold=True)
        
        # Título
        ws.merge_cells('A1:H1')
        ws['A1'] = "Entradas de Inventario"
        ws['A1'].font = title_font
        ws['A1'].alignment = Alignment(horizontal='center')

        # Información de la empresa
        ws['A3'] = "CCD INGENIERIA Y CONSTRUCCIONES S.A.S."
        ws['A4'] = "901092189-5"
        ws['A5'] = "CCD INGENIERIA Y CONSTRUCCIONES"
        
        ws['F3'] = f"Fecha: {fecha_inicio} - {fecha_fin}"
        ws['F4'] = "Dirección: CR 78 A 45 G G 14"
        ws['F5'] = "Teléfono: 3043821361"

        # Encabezados
        headers = ['Numero', 'Fecha y Hora', 'Codigo', 'Producto', 'Cant.', 'Valor', 'Total', 'Observaciones']
        for col, header in enumerate(headers, start=1):
            cell = ws.cell(row=7, column=col, value=header)
            cell.font = header_font
            cell.fill = PatternFill(start_color="4F81BD", end_color="4F81BD", fill_type="solid")
            cell.alignment = Alignment(horizontal='center')

        # Datos
        for row, entrada in enumerate(entradas, start=8):
            ws.cell(row=row, column=1, value=entrada['Numero'])
            ws.cell(row=row, column=2, value=entrada['Fecha'])
            ws.cell(row=row, column=3, value=entrada['IdReferencia'])
            ws.cell(row=row, column=4, value=entrada['Descripcion'])
            ws.cell(row=row, column=5, value=entrada['Cantidad'])
            ws.cell(row=row, column=6, value=entrada['Valor'])
            ws.cell(row=row, column=7, value=entrada['Total'])
            ws.cell(row=row, column=8, value=entrada.get('Observaciones', ''))

        # Ajustar el ancho de las columnas
        for column in ws.columns:
            max_length = 0
            column_letter = get_column_letter(column[0].column)
            for cell in column:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            adjusted_width = (max_length + 2)
            ws.column_dimensions[column_letter].width = adjusted_width

        current_app.logger.info("Archivo Excel creado exitosamente")
        
        output = io.BytesIO()
        wb.save(output)
        output.seek(0)

        response = make_response(output.getvalue())
        response.headers['Content-Disposition'] = 'attachment; filename=Entradas_de_Inventario.xlsx'
        response.headers['Content-type'] = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        response.headers['Access-Control-Allow-Origin'] = '*'
        
        current_app.logger.info("Archivo Excel enviado exitosamente")
        return response

    except ValueError as ve:
        current_app.logger.error(f"Error de validación: {str(ve)}")
        return jsonify({"error": str(ve)}), 400
    except Exception as e:
        current_app.logger.error(f"Error al exportar entradas: {str(e)}")
        current_app.logger.error(traceback.format_exc())
        return jsonify({"error": "Error interno del servidor", "details": str(e), "traceback": traceback.format_exc()}), 500

@app.route('/api/debug_exportar_entradas_excel')
@cross_origin(supports_credentials=True)
def debug_exportar_entradas_excel():
    try:
        fecha_inicio = request.args.get('fecha_inicio')
        fecha_fin = request.args.get('fecha_fin')
        
        current_app.logger.info(f"Iniciando depuración de exportación de entradas: {fecha_inicio} - {fecha_fin}")
        
        # Paso 1: Obtener datos
        entradas = obtener_datos_entradas(fecha_inicio, fecha_fin)
        current_app.logger.info(f"Datos obtenidos: {len(entradas)} entradas")
        
        if not entradas:
            raise ValueError("No hay datos para exportar")

        # Paso 2: Crear workbook
        wb = Workbook()
        ws = wb.active
        ws.title = "Entradas de Inventario"
        current_app.logger.info("Workbook creado")
        
        # Paso 3: Añadir encabezados
        headers = ['Numero', 'Fecha', 'Codigo', 'Producto', 'Cant.', 'Valor', 'Total']
        for col, header in enumerate(headers, start=1):
            ws.cell(row=1, column=col, value=header)
        current_app.logger.info("Encabezados añadidos")
        
        # Paso 4: Añadir datos
        for row, entrada in enumerate(entradas, start=2):
            ws.cell(row=row, column=1, value=entrada['Numero'])
            ws.cell(row=row, column=2, value=entrada['Fecha'])
            ws.cell(row=row, column=3, value=entrada['IdReferencia'])
            ws.cell(row=row, column=4, value=entrada['Descripcion'])
            ws.cell(row=row, column=5, value=entrada['Cantidad'])
            ws.cell(row=row, column=6, value=entrada['Valor'])
            ws.cell(row=row, column=7, value=entrada['Total'])
        current_app.logger.info("Datos añadidos")
        
        # Paso 5: Guardar en BytesIO
        output = io.BytesIO()
        wb.save(output)
        output.seek(0)
        current_app.logger.info("Archivo guardado en BytesIO")
        
        # Paso 6: Enviar archivo
        response = send_file(
            output,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name='Debug_Entradas_de_Inventario.xlsx'
        )
        response.headers["Access-Control-Allow-Origin"] = "*"
        return response

    except ValueError as ve:
        current_app.logger.error(f"Error de validación en depuración: {str(ve)}")
        return jsonify({"error": str(ve)}), 400
    except Exception as e:
        current_app.logger.error(f"Error en depuración de exportación de entradas: {str(e)}")
        current_app.logger.error(traceback.format_exc())
        return jsonify({"error": str(e), "traceback": traceback.format_exc()}), 500

def obtener_datos_entradas(fecha_inicio, fecha_fin):
    # Convertir las fechas de string a objetos datetime
    fecha_inicio = datetime.strptime(fecha_inicio, '%Y-%m-%d').date()
    fecha_fin = datetime.strptime(fecha_fin, '%Y-%m-%d').date()
    
    # Consulta para obtener las entradas
    entradas = db.session.query(
        Entradas1.Numero,
        Entradas1.fecha,
        Entradas2.IdReferencia,
        Entradas2.Descripcion,
        Entradas2.Cantidad,
        Entradas2.Valor,
        (Entradas2.Cantidad * Entradas2.Valor).label('Total'),
        Entradas1.Observaciones
    ).join(
        Entradas2,
        Entradas1.Numero == Entradas2.Numero
    ).filter(
        Entradas1.fecha.between(fecha_inicio, fecha_fin)
    ).all()
    
    # Convertir los resultados a un formato serializable
    return [{
        'Numero': e.Numero,
        'Fecha': e.fecha.strftime('%Y-%m-%d %H:%M:%S'),
        'IdReferencia': e.IdReferencia,
        'Descripcion': e.Descripcion,
        'Cantidad': float(e.Cantidad),
        'Valor': float(e.Valor),
        'Total': float(e.Total),
        'Observaciones': e.Observaciones if e.Observaciones is not None else ''
    } for e in entradas]

def obtener_datos_salidas(fecha_inicio, fecha_fin):
    # Convertir las fechas de string a objetos datetime
    fecha_inicio = datetime.strptime(fecha_inicio, '%Y-%m-%d').date()
    fecha_fin = datetime.strptime(fecha_fin, '%Y-%m-%d').date()
    
    # Consulta para obtener las salidas
    salidas = db.session.query(
        Salidas1.Numero,
        Salidas1.fecha,
        Salidas2.IdReferencia,
        Salidas2.Descripcion,
        Salidas2.Cantidad,
        Salidas2.Valor,
        (Salidas2.Cantidad * Salidas2.Valor).label('Total'),
        Salidas1.Observaciones
    ).join(
        Salidas2,
        Salidas1.Numero == Salidas2.Numero
    ).filter(
        Salidas1.fecha.between(fecha_inicio, fecha_fin)
    ).all()
    
    # Convertir los resultados a un formato serializable
    return [{
        'Numero': s.Numero,
        'Fecha': s.fecha.strftime('%Y-%m-%d %H:%M:%S'),
        'IdReferencia': s.IdReferencia,
        'Descripcion': s.Descripcion,
        'Cantidad': float(s.Cantidad),
        'Valor': float(s.Valor),
        'Total': float(s.Total),
        'Observaciones': s.Observaciones if s.Observaciones is not None else ''
    } for s in salidas]

@app.route('/api/todas_ciudades', methods=['GET'])
def obtener_todas_ciudades():
    ciudades = Ciudades.query.all()
    return jsonify([{
        'IdCiudad': c.IdCiudad,
        'Ciudad': c.Ciudad,
        'IdDepartamento': c.IdDepartamento,
        'idpais': c.idpais,
        'porcreteica': float(c.porcreteica) if c.porcreteica else None,
        'Estado': c.Estado
    } for c in ciudades])

@app.route('/api/ciudades/<string:id>', methods=['PUT'])
def actualizar_ciudad(id):
    ciudad = Ciudades.query.get(id)
    if not ciudad:
        return jsonify({'message': 'Ciudad no encontrada'}), 404
    data = request.json
    ciudad.Ciudad = data['Ciudad']
    ciudad.IdDepartamento = data['IdDepartamento']
    ciudad.idpais = data['idpais']
    ciudad.porcreteica = data.get('porcreteica')
    ciudad.Estado = data['Estado']
    db.session.commit()
    return jsonify({'message': 'Ciudad actualizada exitosamente'})

@app.route('/api/ciudades/<string:id>', methods=['DELETE'])
def eliminar_ciudad(id):
    ciudad = Ciudades.query.get(id)
    if not ciudad:
        return jsonify({'message': 'Ciudad no encontrada'}), 404
    db.session.delete(ciudad)
    db.session.commit()
    return jsonify({'message': 'Ciudad eliminada exitosamente'})

@app.route('/api/departamentos/<string:id_pais>', methods=['GET'])
def obtener_departamentos_por_pais(id_pais):
    departamentos = Departamentos.query.filter_by(IdPais=id_pais).all()
    return jsonify([{
        'IdDepartamento': d.IdDepartamento,
        'Departamento': d.Departamento
    } for d in departamentos])

@app.route('/api/subgrupos', methods=['GET', 'POST'])
def manejar_subgrupos():
    if request.method == 'POST':
        data = request.json
        nuevo_subgrupo = SubGrupos(
            IdSubgrupo=data['IdSubgrupo'],
            Subgrupo=data['Subgrupo'],
            IdGrupo=data['IdGrupo'],
            Estado=data['Estado']
        )
        db.session.add(nuevo_subgrupo)
        db.session.commit()
        return jsonify({'message': 'Subgrupo creado exitosamente'}), 201
    else:
        subgrupos = SubGrupos.query.all()
        return jsonify([{
            'IdSubgrupo': s.IdSubgrupo,
            'Subgrupo': s.Subgrupo,
            'IdGrupo': s.IdGrupo,
            'Estado': s.Estado
        } for s in subgrupos])

@app.route('/api/subgrupos/<string:id>', methods=['DELETE'])
def eliminar_subgrupo(id):
    subgrupo = SubGrupos.query.get(id)
    if subgrupo:
        db.session.delete(subgrupo)
        db.session.commit()
        return jsonify({'message': 'Subgrupo eliminado exitosamente'})
    return jsonify({'message': 'Subgrupo no encontrado'}), 404

@app.route('/api/grupos', methods=['GET', 'POST'])
def manejar_grupos():
    if request.method == 'GET':
        try:
            grupos = Grupo.query.all()
            return jsonify([{
                'codigo': grupo.IdGrupo,
                'descripcion': grupo.Grupo,
                'estado': grupo.Estado,
                'menupos': grupo.menupos
            } for grupo in grupos])
        except Exception as e:
            print("Error al obtener grupos:", str(e))
            return jsonify({'error': 'Error al obtener grupos: ' + str(e)}), 500
    elif request.method == 'POST':
        try:
            data = request.json
            app.logger.info(f"Procesando datos POST: {data}")
            
            grupo_existente = Grupo.query.get(data['codigo'])
            if grupo_existente:
                app.logger.info(f"Actualizando grupo existente: {data['codigo']}")
                grupo_existente.Grupo = data['descripcion']
                grupo_existente.Estado = data['estado']
                grupo_existente.menupos = data['menupos']
                mensaje = 'Grupo actualizado exitosamente'
            else:
                app.logger.info(f"Creando nuevo grupo: {data['codigo']}")
                nuevo_grupo = Grupo(
                    IdGrupo=data['codigo'],
                    Grupo=data['descripcion'],
                    Estado=data['estado'],
                    menupos=data['menupos']
                )
                db.session.add(nuevo_grupo)
                mensaje = 'Grupo creado exitosamente'
            
            db.session.commit()
            app.logger.info("Operación en base de datos completada con éxito")
            return jsonify({'success': True, 'message': mensaje}), 200
        except IntegrityError as e:
            db.session.rollback()
            print("Error de integridad:", str(e))
            return jsonify({'error': 'El código de grupo ya existe. Por favor, use un código diferente.'}), 400
        except Exception as e:
            db.session.rollback()
            app.logger.error(f"Error al procesar la solicitud: {str(e)}")
            return jsonify({'error': 'Error inesperado: ' + str(e)}), 500
        
    app.logger.warning(f"Método no soportado: {request.method}")
    return jsonify({'error': 'Método no soportado'}), 405

@app.route('/api/bodegas/<string:id>', methods=['DELETE'])
def eliminar_bodega(id):
    try:
        bodega = Bodegas.query.get(id)
        if bodega:
            db.session.delete(bodega)
            db.session.commit()
            return jsonify({'success': True, 'message': 'Bodega eliminada exitosamente'})
        else:
            return jsonify({'success': False, 'message': 'Bodega no encontrada'}), 404
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'message': str(e)}), 500

def verificar_credenciales(id_usuario, contraseña):
    usuario = Usuarios.query.filter_by(IdUsuario=id_usuario).first()
    if usuario and usuario.Contraseña == contraseña:
        return jsonify({"success": True, "message": "Inicio de sesión exitoso"})
    return jsonify({"success": False, "message": "Credenciales inválidas"})

@app.route('/registro')
def registro():
    return render_template('pideLicencia.html')

@app.route('/restContrasenha')
def restContrasenha():
    return render_template('restContrasenha.html')

@app.route('/api/validar_usuario', methods=['GET'])
def validar_usuario():
    id_usuario = request.args.get('idUsuario')
    usuario = Usuarios.query.get(id_usuario)
    return jsonify({'existe': usuario is not None})

@app.route('/api/actualizar_contrasena', methods=['POST'])
def actualizar_contrasena():
    try:
        data = request.json
        id_usuario = data.get('idUsuario')
        nueva_contrasena = data.get('contrasena')

        if not id_usuario or not nueva_contrasena:
            return jsonify({'success': False, 'message': 'Datos incompletos'}), 400

        usuario = Usuarios.query.get(id_usuario)
        if usuario:
            # Guardar la contraseña exactamente como la ingresó el usuario
            usuario.Contraseña = nueva_contrasena
            db.session.commit()

            return jsonify({
                'success': True, 
                'message': 'Contraseña actualizada exitosamente',
                'redirect_url': url_for('index')
            })
        else:
            return jsonify({'success': False, 'message': 'Usuario no encontrado'}), 404
    except Exception as e:
        db.session.rollback()
        app.logger.error(f"Error al actualizar la contraseña: {str(e)}")
        return jsonify({'success': False, 'message': 'Error interno del servidor'}), 500

def login_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'user_id' not in session:
            return redirect(url_for('login', next=request.url))
        return f(*args, **kwargs)
    return decorated_function

@app.route('/')
def index():
    if 'user_id' in session:
        return redirect(url_for('inicio'))
    return redirect(url_for('login'))

@app.route('/inicio')
@login_required
def inicio():
    return render_template('index.html')

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'GET':
        return render_template('login.html')
    return "Método no permitido", 405

@app.route('/api/login', methods=['POST'])
def api_login():
    try:
        data = request.json
        id_usuario = data.get('username')
        contraseña = data.get('password')

        if not id_usuario or not contraseña:
            return jsonify({"success": False, "message": "Usuario y contraseña son requeridos"}), 400

        usuario = Usuarios.query.filter_by(IdUsuario=id_usuario).first()
        
        if usuario and usuario.Contraseña == contraseña:  # Comparación directa de contraseñas
            # Obtener permisos del usuario
            permisos = db.session.query(Permisos.NombrePermiso).join(UsuariosPermisos).filter(UsuariosPermisos.IdUsuario == usuario.IdUsuario).all()
            permisos = [p[0] for p in permisos]

            # Guardar información en la sesión
            session['user_id'] = usuario.IdUsuario
            session['nivel_acceso'] = usuario.NivelAcceso

            return jsonify({
                "success": True, 
                "message": f"Bienvenido {usuario.IdUsuario}",
                "user": {
                    "IdUsuario": usuario.IdUsuario,
                    "Descripcion": usuario.Descripcion,
                    "NivelAcceso": usuario.NivelAcceso,
                    "Permisos": permisos
                },
                "redirect_url": url_for('inicio')
            })
        else:
            return jsonify({"success": False, "message": "Credenciales inválidas"}), 401
    except Exception as e:
        app.logger.error(f"Error en el servidor: {str(e)}")
        return jsonify({"success": False, "message": "Error interno del servidor"}), 500

@app.route('/api/usuarios', methods=['POST'])
@login_required
def crear_usuario():
    try:
        data = request.json
        if not data.get('Contraseña'):
            return jsonify({"success": False, "message": "La contraseña es requerida"}), 400

        nuevo_usuario = Usuarios(
            IdUsuario=data['IdUsuario'],
            Contraseña=data['Contraseña'],  # Guardamos la contraseña tal cual, sin aplicar hash
            Descripcion=data['Descripcion'],
            Estado=data['Estado'],
            NivelAcceso=data['NivelAcceso']
        )
        db.session.add(nuevo_usuario)
        db.session.commit()

        return jsonify({"success": True, "usuario": {
            "IdUsuario": nuevo_usuario.IdUsuario,
            "Descripcion": nuevo_usuario.Descripcion,
            "NivelAcceso": nuevo_usuario.NivelAcceso
        }}), 201

    except Exception as e:
        db.session.rollback()
        app.logger.error(f"Error al crear usuario: {str(e)}")
        return jsonify({"success": False, "message": f"Error al crear usuario: {str(e)}"}), 500

@app.route('/api/usuarios/<string:user_id>', methods=['DELETE'])
@login_required
def eliminar_usuario(user_id):
    if session.get('nivel_acceso') != 100:
        return jsonify({'success': False, 'message': 'No tienes permiso para eliminar usuarios'}), 403

    try:
        usuario = Usuarios.query.get(user_id)
        if usuario:
            db.session.delete(usuario)
            db.session.commit()
            return jsonify({'success': True, 'message': 'Usuario eliminado exitosamente'})
        else:
            return jsonify({'success': False, 'message': 'Usuario no encontrado'}), 404
    except Exception as e:
        db.session.rollback()
        app.logger.error(f"Error al eliminar usuario: {str(e)}")
        return jsonify({'success': False, 'message': 'Error al eliminar usuario'}), 500

@app.route('/api/usuarios/<string:user_id>', methods=['PUT'])
@login_required
def actualizar_usuario(user_id):
    if session.get('nivel_acceso') != 100:
        return jsonify({'success': False, 'message': 'No tienes permiso para actualizar usuarios'}), 403

    try:
        usuario = Usuarios.query.get(user_id)
        if not usuario:
            return jsonify({'success': False, 'message': 'Usuario no encontrado'}), 404

        data = request.json
        usuario.Descripcion = data.get('Descripcion', usuario.Descripcion)
        usuario.Estado = data.get('Estado', usuario.Estado)
        usuario.NivelAcceso = data.get('NivelAcceso', usuario.NivelAcceso)

        # ADVERTENCIA: Esto almacena la contraseña en texto plano
        if 'Contraseña' in data and data['Contraseña']:
            usuario.Contraseña = data['Contraseña']

        db.session.commit()
        return jsonify({'success': True, 'message': 'Usuario actualizado exitosamente'})
    except Exception as e:
        db.session.rollback()
        app.logger.error(f"Error al actualizar usuario: {str(e)}")
        return jsonify({'success': False, 'message': f'Error al actualizar usuario: {str(e)}'}), 500

@app.route('/api/usuarios', methods=['GET'])
@login_required
def obtener_usuarios():
    try:
        usuarios = Usuarios.query.all()
        resultado = [{
            'IdUsuario': u.IdUsuario,
            'Descripcion': u.Descripcion,
            'NivelAcceso': u.NivelAcceso,
            'Estado': u.Estado
        } for u in usuarios]
        app.logger.info(f"Usuarios obtenidos: {resultado}")
        return jsonify(resultado)
    except Exception as e:
        app.logger.error(f"Error al obtener usuarios: {str(e)}")
        return jsonify({'error': 'Error al obtener usuarios'}), 500

def admin_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        logger.debug(f"Verificando permisos de administrador para {current_user}")
        if not current_user.is_authenticated:
            logger.warning("Usuario no autenticado intentando acceder a ruta protegida")
            return jsonify({"success": False, "message": "No autenticado"}), 401
        if current_user.NivelAcceso != 100:
            logger.warning(f"Usuario {current_user.IdUsuario} sin permisos de administrador intentando acceder a ruta protegida")
            return jsonify({"success": False, "message": "Acceso no autorizado"}), 403
        return f(*args, **kwargs)
    return decorated_function

@app.route('/api/usuarios', methods=['GET', 'POST'])
@login_required
@admin_required
def manejar_usuarios():
    try:
        if request.method == 'GET':
            logger.info("Obteniendo lista de usuarios")
            usuarios = Usuarios.query.all()
            return jsonify([{
                'IdUsuario': u.IdUsuario,
                'Descripcion': u.Descripcion,
                'NivelAcceso': u.NivelAcceso,
                'Estado': u.Estado
            } for u in usuarios])
        elif request.method == 'POST':
            logger.info("Creando nuevo usuario")
            data = request.json
            if not data.get('Contraseña'):
                return jsonify({"success": False, "message": "La contraseña es requerida"}), 400

            nuevo_usuario = Usuarios(
                IdUsuario=data['IdUsuario'],
                Contraseña=generate_password_hash(data['Contraseña']),
                Descripcion=data['Descripcion'],
                Estado=data['Estado'],
                NivelAcceso=data['NivelAcceso']
            )
            db.session.add(nuevo_usuario)
            db.session.commit()

            return jsonify({"success": True, "usuario": {
                "IdUsuario": nuevo_usuario.IdUsuario,
                "Descripcion": nuevo_usuario.Descripcion,
                "NivelAcceso": nuevo_usuario.NivelAcceso,
                "Estado": nuevo_usuario.Estado
            }}), 201
    except Exception as e:
        logger.error(f"Error en manejar_usuarios: {str(e)}")
        return jsonify({"success": False, "message": "Error interno del servidor"}), 500

@app.route('/api/usuarios/<string:user_id>', methods=['GET', 'PUT', 'DELETE'])
@login_required
def manejar_usuario_individual(user_id):
    try:
        logger.info(f"Manejando solicitud para usuario {user_id}")
        usuario = Usuarios.query.get(user_id)
        if not usuario:
            logger.warning(f"Usuario {user_id} no encontrado")
            return jsonify({"success": False, "message": "Usuario no encontrado"}), 404

        if request.method == 'GET':
            logger.info(f"Obteniendo datos del usuario {user_id}")
            return jsonify({
                "success": True,
                "usuario": {
                    "IdUsuario": usuario.IdUsuario,
                    "Descripcion": usuario.Descripcion,
                    "NivelAcceso": usuario.NivelAcceso,
                    "Estado": usuario.Estado
                }
            })

        elif request.method == 'PUT':
            logger.info(f"Actualizando usuario {user_id}")
            data = request.json
            usuario.Descripcion = data.get('Descripcion', usuario.Descripcion)
            usuario.Estado = data.get('Estado', usuario.Estado)
            usuario.NivelAcceso = data.get('NivelAcceso', usuario.NivelAcceso)

            if 'Contraseña' in data and data['Contraseña']:
                usuario.Contraseña = generate_password_hash(data['Contraseña'])

            db.session.commit()
            return jsonify({'success': True, 'message': 'Usuario actualizado exitosamente'})

        elif request.method == 'DELETE':
            logger.info(f"Eliminando usuario {user_id}")
            db.session.delete(usuario)
            db.session.commit()
            return jsonify({'success': True, 'message': 'Usuario eliminado exitosamente'})

    except Exception as e:
        logger.error(f"Error en manejar_usuario_individual: {str(e)}", exc_info=True)
        db.session.rollback()
        return jsonify({"success": False, "message": f"Error interno del servidor: {str(e)}"}), 500


@app.route('/api/permisos', methods=['GET'])
@login_required
@admin_required
def obtener_permisos():
    try:
        logger.info("Obteniendo lista de permisos")
        permisos = UsuariosPermisos.query.all()
        return jsonify([{
            'id': p.IdPermiso,
            'nombre': p.NombrePermiso
        } for p in permisos])
    except Exception as e:
        logger.error(f"Error en obtener_permisos: {str(e)}")
        return jsonify({"success": False, "message": "Error interno del servidor"}), 500

@app.errorhandler(500)
def internal_server_error(e):
    logger.error(f"Error 500: {str(e)}")
    return jsonify(error="Error interno del servidor"), 500

# Asegúrate de que esta función esté definida en tu aplicación
@app.errorhandler(405)
def method_not_allowed(e):
    return jsonify(error=str(e)), 405

@app.route('/logout')
def logout():
    # Eliminar todos los datos de la sesión
    session.clear()
    # Redirigir al usuario a la página de login
    return redirect(url_for('login'))

def login_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'user_id' not in session:
            return redirect(url_for('login', next=request.url))
        return f(*args, **kwargs)
    return decorated_function

@app.after_request
def add_header(response):
    response.headers['Cache-Control'] = 'no-store, no-cache, must-revalidate, post-check=0, pre-check=0, max-age=0'
    response.headers['Pragma'] = 'no-cache'
    response.headers['Expires'] = '-1'
    return response

@app.route('/api/bancos', methods=['GET'])
def obtener_bancos():
    try:
        bancos = Bancos.query.all()
        return jsonify([{
            'IdBanco': banco.IdBanco,
            'Banco': banco.Banco,
            'Estado': banco.Estado,
            'cuenta': banco.cuenta,
            'tipo': banco.tipo,
            'cuentacontable': banco.cuentacontable,
            'chequeactual': banco.chequeactual,
            'chequefinal': banco.chequefinal,
            'escheque': banco.escheque,
            'tipocuenta': banco.tipocuenta
        } for banco in bancos])
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/bancos', methods=['POST'])
def crear_actualizar_banco():
    data = request.json
    try:
        banco = Bancos.query.get(data['IdBanco'])
        if banco:
            # Actualizar banco existente
            for key, value in data.items():
                setattr(banco, key, value)
        else:
            # Crear nuevo banco
            banco = Bancos(**data)
            db.session.add(banco)
        
        db.session.commit()
        return jsonify({'message': 'Banco guardado exitosamente', 'banco': data}), 200
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500

@app.route('/api/bancos/<string:id>', methods=['DELETE'])
def eliminar_banco(id):
    try:
        banco = Bancos.query.get(id)
        if banco:
            db.session.delete(banco)
            db.session.commit()
            return jsonify({'message': 'Banco eliminado exitosamente'}), 200
        else:
            return jsonify({'message': 'Banco no encontrado'}), 404
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500
    
@app.route('/api/consecutivos_salidas_inventario', methods=['GET'])
def obtener_consecutivos_salidas_inventario():
    try:
        # Suponiendo que 'SAL' es el código para Salidas de Inventario
        consecutivos = Consecutivos.query.filter_by(Formulario='SAL', Estado=True).all()
        return jsonify([{
            'IdConsecutivo': c.IdConsecutivo,
            'Descripcion': c.Consecutivo,
            'Prefijo': c.Prefijo,
            'Actual': c.Actual
        } for c in consecutivos])
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/consecutivos_compra', methods=['GET'])
def obtener_consecutivos_compra():
    try:
        # Asumiendo que 'CM' es el código para Compras de Mercancía
        consecutivos = Consecutivos.query.filter_by(Formulario='CM', Estado=True).all()
        return jsonify([{
            'IdConsecutivo': c.IdConsecutivo,
            'Descripcion': c.Consecutivo,
            'Prefijo': c.Prefijo,
            'Actual': c.Actual
        } for c in consecutivos])
    except Exception as e:
        app.logger.error(f"Error al obtener consecutivos de compra: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/ultimo_consecutivo_compras', methods=['GET'])
def obtener_ultimo_consecutivo_compras():
    try:
        consecutivo = Consecutivos.query.filter_by(Formulario='CM').order_by(Consecutivos.Actual.desc()).first()
        if consecutivo:
            ultimo_consecutivo = f"{consecutivo.Prefijo}{consecutivo.Actual.zfill(2)}"
            return jsonify({'success': True, 'ultimoConsecutivo': ultimo_consecutivo})
        else:
            return jsonify({'success': False, 'message': 'No se encontró el consecutivo para Compras'})
    except Exception as e:
        app.logger.error(f"Error al obtener el último consecutivo de compras: {str(e)}")
        return jsonify({'success': False, 'message': str(e)}), 500
    
@app.route('/api/ciudades/<string:id_departamento>', methods=['GET'])
def obtener_ciudades(id_departamento):
    ciudades = Ciudades.query.filter_by(IdDepartamento=id_departamento).all()
    return jsonify([{
        'IdCiudad': c.IdCiudad,
        'Ciudad': c.Ciudad
    } for c in ciudades])
    
@app.route('/api/consecutivos_ordenes_compra', methods=['GET'])
def obtener_consecutivos_ordenes_compra():
    try:
        # Suponiendo que 'OC' es el código para Órdenes de Compra
        consecutivos = Consecutivos.query.filter_by(Formulario='OC', Estado=True).all()
        return jsonify([{
            'IdConsecutivo': c.IdConsecutivo,
            'Descripcion': c.Consecutivo,
            'Prefijo': c.Prefijo,
            'Actual': c.Actual
        } for c in consecutivos])
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/consecutivos_traslados_inventario', methods=['GET'])
def obtener_consecutivos_traslados_inventario():
    try:
        # Suponiendo que 'TB' es el código para Traslados de Inventario
        consecutivos = Consecutivos.query.filter_by(Formulario='TB', Estado=True).all()
        return jsonify([{
            'IdConsecutivo': c.IdConsecutivo,
            'Descripcion': c.Consecutivo,
            'Prefijo': c.Prefijo,
            'Actual': c.Actual
        } for c in consecutivos])
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/guardar_traslado', methods=['POST'])
def guardar_traslado():
    data = request.json
    traslado1_data = data.get('traslado1', {})
    traslados2_data = data.get('traslados2', [])

    try:
        # Verificar y limpiar referencias
        referencias_invalidas = []
        for traslado2 in traslados2_data:
            id_referencia = traslado2.get('IdReferencia')
            if id_referencia:
                # Limpiar el ID de referencia
                id_referencia_limpio = ''.join(c for c in id_referencia if c.isalnum() or c in '-_')
                traslado2['IdReferencia'] = id_referencia_limpio
                if not Referencia.query.get(id_referencia_limpio):
                    referencias_invalidas.append(id_referencia_limpio)
            else:
                referencias_invalidas.append('ID vacío')

        if referencias_invalidas:
            return jsonify({
                'success': False,
                'message': f'Referencias inválidas encontradas: {", ".join(referencias_invalidas)}'
            }), 400

        nuevo_traslado1 = Traslados1(
            Numero=traslado1_data.get('Numero'),
            Mes=traslado1_data.get('Mes'),
            Anulado=traslado1_data.get('Anulado', False),
            IdBodegaOrigen=traslado1_data.get('IdBodegaOrigen'),
            IdBodegaDestino=traslado1_data.get('IdBodegaDestino'),
            Observaciones=traslado1_data.get('Observaciones'),
            FechaCreacion=datetime.fromisoformat(traslado1_data.get('FechaCreacion')),
            IdUsuario=traslado1_data.get('IdUsuario'),
            IdConsecutivo=traslado1_data.get('IdConsecutivo'),
            fecha=datetime.strptime(traslado1_data.get('fecha'), '%Y-%m-%d'),
            subtotal=Decimal(str(traslado1_data.get('subtotal', '0'))),
            total_iva=Decimal(str(traslado1_data.get('total_iva', '0'))),
            total_impoconsumo=Decimal(str(traslado1_data.get('total_impoconsumo', '0'))),
            total_ipc=Decimal(str(traslado1_data.get('total_ipc', '0'))),
            total_ibua=Decimal(str(traslado1_data.get('total_ibua', '0'))),
            total_icui=Decimal(str(traslado1_data.get('total_icui', '0'))),
            total=Decimal(str(traslado1_data.get('total', '0')))
        )
        db.session.add(nuevo_traslado1)

        mes_actual = datetime.now().strftime('%Y%m')
        
        for traslado2 in traslados2_data:
            cantidad_trasladada = Decimal(str(traslado2.get('Cantidad', '0')))
            id_referencia = traslado2['IdReferencia']
            
            # Actualizar saldo en bodega origen
            actualizar_saldo_bodega(
                traslado1_data['IdBodegaOrigen'], 
                id_referencia, 
                mes_actual, 
                -cantidad_trasladada
            )
            
            # Verificar si el producto existe en la bodega destino, si no, crearlo
            referencia_destino = obtener_o_crear_referencia(
                traslado1_data['IdBodegaDestino'],
                id_referencia,
                traslado2
            )
            
            # Actualizar saldo en bodega destino
            actualizar_saldo_bodega(
                traslado1_data['IdBodegaDestino'], 
                referencia_destino.IdReferencia, 
                mes_actual, 
                cantidad_trasladada
            )
            
            nuevo_traslado2 = Traslados2(
                ID=traslado2.get('ID'),
                Numero=traslado2.get('Numero'),
                IdReferencia=id_referencia,
                Descripcion=traslado2.get('Descripcion'),
                Cantidad=cantidad_trasladada,
                Valor=Decimal(str(traslado2.get('Valor', '0'))),
                IVA=Decimal(str(traslado2.get('IVA', '0'))),
                idunidad=traslado2.get('idunidad'),
                impoconsumo=Decimal(str(traslado2.get('impoconsumo', '0'))),
                ipc=Decimal(str(traslado2.get('ipc', '0'))),
                imp_ibua=Decimal(str(traslado2.get('imp_ibua', '0'))),
                imp_icui=Decimal(str(traslado2.get('imp_icui', '0')))
            )
            db.session.add(nuevo_traslado2)

        db.session.commit()
        return jsonify({'success': True, 'message': 'Traslado guardado con éxito'})
    except Exception as e:
        db.session.rollback()
        app.logger.error(f"Error al guardar el traslado: {str(e)}")
        return jsonify({'success': False, 'message': f'Error al guardar el traslado: {str(e)}'}), 500

def actualizar_saldo_bodega(id_bodega, id_referencia, mes, cantidad):
    try:
        saldo = SaldosBodega.query.filter_by(
            IdBodega=id_bodega,
            IdReferencia=id_referencia,
            Mes=mes
        ).first()
        
        if saldo:
            saldo.Saldo += cantidad
            if cantidad > 0:
                saldo.Entradas += cantidad
            else:
                saldo.Salidas += abs(cantidad)
        else:
            nuevo_saldo = SaldosBodega(
                IdBodega=id_bodega,
                IdReferencia=id_referencia,
                Mes=mes,
                Saldo=cantidad,
                Entradas=max(cantidad, 0),
                Salidas=max(-cantidad, 0),
                SaldoInicial=0
            )
            db.session.add(nuevo_saldo)
    except Exception as e:
        app.logger.error(f"Error al actualizar saldo de bodega: {str(e)}")
        raise

def obtener_o_crear_referencia(id_bodega_destino, id_referencia_origen, datos_producto):
    referencia_destino = Referencia.query.filter_by(
        IdReferencia=id_referencia_origen, 
        idbodega=id_bodega_destino
    ).first()

    if not referencia_destino:
        # El producto no existe en la bodega destino, creamos uno nuevo
        referencia_origen = Referencia.query.get(id_referencia_origen)
        if referencia_origen:
            nuevo_id_referencia = generar_nuevo_id_referencia(referencia_origen.IdGrupo)
            referencia_destino = Referencia(
                IdReferencia=nuevo_id_referencia,
                Referencia=referencia_origen.Referencia,
                IdGrupo=referencia_origen.IdGrupo,
                IdUnidad=referencia_origen.IdUnidad,
                idbodega=id_bodega_destino,
                PrecioVenta1=referencia_origen.PrecioVenta1,
                IVA=referencia_origen.IVA,
                Costo=referencia_origen.Costo,
                Marca=referencia_origen.Marca,
                EstadoProducto=referencia_origen.EstadoProducto,
                Estado=referencia_origen.Estado,
                Tipo=referencia_origen.Tipo,
                ManejaInventario=referencia_origen.ManejaInventario,
                # Añade aquí otros campos necesarios
            )
            db.session.add(referencia_destino)
        else:
            raise ValueError(f"Referencia de origen no encontrada: {id_referencia_origen}")
    
    return referencia_destino

def generar_nuevo_id_referencia(id_grupo):
    grupo = Grupo.query.get(id_grupo)
    if grupo:
        grupo.ultimoCodigo += 1
        return f"{id_grupo}{grupo.ultimoCodigo:02d}"
    else:
        raise ValueError(f"Grupo no encontrado: {id_grupo}")

@app.route('/api/ultimo_consecutivo_traslados', methods=['GET'])
def ultimo_consecutivo_traslados():
    try:
        consecutivo = Consecutivos.query.filter_by(Formulario='TB').first()
        if consecutivo:
            ultimo_consecutivo = f"{consecutivo.Prefijo}{consecutivo.Actual.zfill(2)}"
            return jsonify({'success': True, 'ultimoConsecutivo': ultimo_consecutivo})
        else:
            return jsonify({'success': False, 'message': 'No se encontró el consecutivo para Traslados de Inventario'})
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)})

@app.route('/api/actualizar_consecutivo_traslados', methods=['POST'])
def actualizar_consecutivo_traslados():
    try:
        consecutivo = Consecutivos.query.filter_by(Formulario='TB').first()
        if consecutivo:
            consecutivo.Actual = str(int(consecutivo.Actual) + 1).zfill(2)
            db.session.commit()
            nuevo_consecutivo = f"{consecutivo.Prefijo}{consecutivo.Actual}"
            return jsonify({'success': True, 'nuevoConsecutivo': nuevo_consecutivo})
        else:
            return jsonify({'success': False, 'message': 'No se encontró el consecutivo para Traslados de Inventario'})
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'message': str(e)})

@app.route('/api/referencias/<string:id>', methods=['PUT'])
def actualizar_referencia(id):
    try:
        data = request.json
        referencia = Referencia.query.get(id)
        
        if not referencia:
            return jsonify({'message': 'Referencia no encontrada'}), 404

        # Actualizar los campos de la referencia
        for key, value in data.items():
            setattr(referencia, key, value)

        db.session.commit()
        return jsonify({'message': 'Referencia actualizada exitosamente'}), 200
    except IntegrityError as e:
        db.session.rollback()
        return jsonify({'message': f'Error de integridad en la base de datos: {str(e)}'}), 400
    except Exception as e:
        db.session.rollback()
        return jsonify({'message': f'Error al actualizar la referencia: {str(e)}'}), 500

@app.route('/api/guardar_compra', methods=['POST'])
def guardar_compra():
    data = request.json
    compra1_data = data.get('compra1', {})
    compras2_data = data.get('compras2', [])

    try:
        # Verificar y limpiar referencias
        referencias_invalidas = []
        for compra2 in compras2_data:
            id_referencia = compra2.get('IdReferencia')
            if id_referencia:
                id_referencia_limpio = ''.join(c for c in id_referencia if c.isalnum() or c in '-_')
                compra2['IdReferencia'] = id_referencia_limpio
                referencia = Referencia.query.get(id_referencia_limpio)
                if not referencia:
                    referencias_invalidas.append(id_referencia_limpio)
                elif not compra2.get('IdGrupo'):
                    compra2['IdGrupo'] = referencia.IdGrupo
            else:
                referencias_invalidas.append('ID vacío')

            if not compra2.get('IdGrupo'):
                referencias_invalidas.append(f'{id_referencia_limpio} (falta IdGrupo)')

        if referencias_invalidas:
            return jsonify({
                'success': False,
                'message': f'Referencias inválidas o incompletas encontradas: {", ".join(referencias_invalidas)}'
            }), 400

        if not compra1_data.get('IdBodega'):
            return jsonify({
                'success': False,
                'message': 'El campo IdBodega es requerido para la compra'
            }), 400

        nueva_compra1 = Compras1(
            Numero=compra1_data.get('Numero'),
            Mes=compra1_data.get('Mes'),
            Anulado=compra1_data.get('Anulado', False),
            Fecha=datetime.strptime(compra1_data.get('Fecha'), '%Y-%m-%d') if compra1_data.get('Fecha') else None,
            FechaCreacion=datetime.now(),
            fechamodificacion=datetime.now(),
            Observaciones=compra1_data.get('Observaciones'),
            IdUsuario=compra1_data.get('IdUsuario'),
            IdBodega=compra1_data.get('IdBodega'),
            Nit=compra1_data.get('Nit'),
            NumFactura=compra1_data.get('NumFactura'),
            IdCentroCosto=compra1_data.get('IdCentroCosto'),
            IdConsecutivo=compra1_data.get('IdConsecutivo'),
            idimpuesto=compra1_data.get('idimpuesto'),
            descuento=safe_decimal(compra1_data.get('descuento')),
            porcretefuente=safe_decimal(compra1_data.get('porcretefuente')),
            retefuente=safe_decimal(compra1_data.get('retefuente')),
            reteica=safe_decimal(compra1_data.get('reteica')),
            reteiva=safe_decimal(compra1_data.get('reteiva')),
            total=safe_decimal(compra1_data.get('total')),
            totaliva=safe_decimal(compra1_data.get('totaliva')),
            documento1=compra1_data.get('documento1'),
            cantidadmetaldisponible=safe_decimal(compra1_data.get('cantidadmetaldisponible')),
            totalcantidadmetal=safe_decimal(compra1_data.get('totalcantidadmetal')),
            topemaximo=safe_decimal(compra1_data.get('topemaximo')),
            acumulado=safe_decimal(compra1_data.get('acumulado')),
            totalcompras=safe_decimal(compra1_data.get('totalcompras')),
            flete=safe_decimal(compra1_data.get('flete')),
            tipoempresa=compra1_data.get('tipoempresa'),
            tipoproveedor=compra1_data.get('tipoproveedor'),
            calculareteiva=compra1_data.get('calculareteiva', False),
            calculareteica=compra1_data.get('calculareteica', False),
            porcuotas=compra1_data.get('porcuotas', False),
            subtotal=safe_decimal(compra1_data.get('subtotal')),
            totaldescuento=safe_decimal(compra1_data.get('totaldescuento')),
            valorindustriacomercio=safe_decimal(compra1_data.get('valorindustriacomercio')),
            transmitido=compra1_data.get('transmitido', False),
            totalipc=safe_decimal(compra1_data.get('totalipc')),
            total_ibua=safe_decimal(compra1_data.get('total_ibua')),
            total_icui=safe_decimal(compra1_data.get('total_icui'))
        )
        db.session.add(nueva_compra1)

        mes_actual = datetime.now().strftime('%Y%m')

        for compra2 in compras2_data:
            cantidad_comprada = Decimal(str(compra2.get('Cantidad', '0')))
            id_referencia_original = compra2['IdReferencia']

            # Verificar si el producto existe en la bodega, si no, crearlo
            referencia = obtener_o_crear_referencia_compra(
                compra1_data['IdBodega'],
                id_referencia_original,
                compra2
            )

            # Actualizar saldo en bodega
            actualizar_saldo_bodega_compra(
                compra1_data['IdBodega'], 
                referencia.IdReferencia,
                mes_actual, 
                cantidad_comprada
            )

            nueva_compra2 = Compras2(
                ID=f"{nueva_compra1.Numero}_{uuid.uuid4().hex[:8]}",
                Numero=nueva_compra1.Numero,
                IdReferencia=referencia.IdReferencia,
                Descripcion=compra2.get('Descripcion'),
                Cantidad=safe_decimal(compra2.get('Cantidad')),
                Valor=safe_decimal(compra2.get('Valor')),
                IVA=safe_decimal(compra2.get('IVA')),
                Descuento=safe_decimal(compra2.get('Descuento')),
                NumOrdenCompra=compra2.get('NumOrdenCompra'),
                NumEntradaCia=compra2.get('NumEntradaCia'),
                idfuente=compra2.get('idfuente'),
                IdCentroCosto=compra2.get('IdCentroCosto'),
                precioventa1=safe_decimal(compra2.get('precioventa1')),
                descuentounitario=safe_decimal(compra2.get('descuentounitario')),
                margen=safe_decimal(compra2.get('margen')),
                ley=safe_decimal(compra2.get('ley')),
                peso=safe_decimal(compra2.get('peso')),
                tipo=compra2.get('tipo'),
                lote=compra2.get('lote'),
                idunidad=compra2.get('idunidad'),
                ipc=safe_decimal(compra2.get('ipc')),
                imp_ibua=safe_decimal(compra2.get('imp_ibua')),
                imp_icui=safe_decimal(compra2.get('imp_icui'))
            )
            db.session.add(nueva_compra2)

        db.session.commit()
        return jsonify({'success': True, 'message': 'Compra guardada con éxito', 'numeroCompra': nueva_compra1.Numero})
    except SQLAlchemyError as e:
        db.session.rollback()
        app.logger.error(f"Error de base de datos al guardar la compra: {str(e)}")
        return jsonify({'success': False, 'message': f'Error de base de datos al guardar la compra: {str(e)}'}), 500
    except Exception as e:
        db.session.rollback()
        app.logger.error(f"Error inesperado al guardar la compra: {str(e)}")
        return jsonify({'success': False, 'message': f'Error inesperado al guardar la compra: {str(e)}'}), 500

def obtener_o_crear_referencia_compra(id_bodega, id_referencia, datos_producto):
    # Primero, intentamos encontrar la referencia exacta en la bodega especificada
    referencia = Referencia.query.filter_by(
        IdReferencia=id_referencia, 
        idbodega=id_bodega
    ).first()

    if referencia:
        # Si la referencia existe en esta bodega, solo actualizamos los datos necesarios
        referencia.Costo = safe_decimal(datos_producto.get('Valor', referencia.Costo))
        referencia.PrecioVenta1 = safe_decimal(datos_producto.get('precioventa1', referencia.PrecioVenta1))
        return referencia

    # Si no existe en esta bodega, buscamos si existe en otra bodega
    referencia_existente = Referencia.query.filter_by(IdReferencia=id_referencia).first()

    if referencia_existente:
        # Si existe en otra bodega, creamos una nueva referencia para esta bodega
        nuevo_id_referencia = generar_nuevo_id_referencia(referencia_existente.IdGrupo)
    else:
        # Si no existe en ninguna bodega, usamos el id_referencia original
        nuevo_id_referencia = id_referencia

    # Creamos la nueva referencia
    nueva_referencia = Referencia(
        IdReferencia=nuevo_id_referencia,
        Referencia=datos_producto.get('Descripcion', referencia_existente.Referencia if referencia_existente else ''),
        IdGrupo=datos_producto.get('IdGrupo', referencia_existente.IdGrupo if referencia_existente else None),
        IdUnidad=datos_producto.get('idunidad', referencia_existente.IdUnidad if referencia_existente else ''),
        idbodega=id_bodega,
        PrecioVenta1=safe_decimal(datos_producto.get('precioventa1')),
        IVA=safe_decimal(datos_producto.get('IVA')),
        Costo=safe_decimal(datos_producto.get('Valor')),
        Estado=True,
        Tipo=False,
        ManejaInventario=True,
        FechaCreacion=datetime.now(),
        StockMinimo=0,
        StockMaximo=0,
        SaldoAntesInv=0,
        Insumo=False,
        costoreal=0
    )

    db.session.add(nueva_referencia)
    return nueva_referencia

def actualizar_saldo_bodega_compra(id_bodega, id_referencia, mes, cantidad):
    try:
        saldo = SaldosBodega.query.filter_by(
            IdBodega=id_bodega,
            IdReferencia=id_referencia,
            Mes=mes
        ).first()

        if saldo:
            saldo.Saldo += cantidad
            saldo.Compras += cantidad
        else:
            nuevo_saldo = SaldosBodega(
                IdBodega=id_bodega,
                IdReferencia=id_referencia,
                Mes=mes,
                Saldo=cantidad,
                Compras=cantidad,
                SaldoInicial=0,
                Entradas=0,
                Salidas=0,
                Ventas=0
            )
            db.session.add(nuevo_saldo)
    except Exception as e:
        app.logger.error(f"Error al actualizar saldo de bodega: {str(e)}")
        raise

def safe_decimal(value, default=Decimal('0')):
    if value is None:
        return default
    try:
        return Decimal(str(value))
    except (InvalidOperation, ValueError, TypeError):
        return default

@app.route('/api/actualizar_consecutivo_compras', methods=['POST'])
def actualizar_consecutivo_compras():
    try:
        consecutivo = Consecutivos.query.filter_by(Formulario='CM').first()
        if consecutivo:
            consecutivo.Actual = str(int(consecutivo.Actual) + 1).zfill(2)
            db.session.commit()
            nuevo_consecutivo = f"{consecutivo.Prefijo}{consecutivo.Actual}"
            return jsonify({'success': True, 'nuevoConsecutivo': nuevo_consecutivo})
        else:
            return jsonify({'success': False, 'message': 'No se encontró el consecutivo para Compras de Mercancía'})
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'message': str(e)})

@app.route('/api/compras_costos_informes', methods=['GET'])
@cross_origin()
def get_compras_costos_informes():
    app.logger.info('Función get_compras_costos_informes llamada')
    fecha_inicio = request.args.get('fecha_inicio')
    fecha_fin = request.args.get('fecha_fin')
   
    app.logger.info(f'Fechas recibidas: inicio={fecha_inicio}, fin={fecha_fin}')
    if not fecha_inicio or not fecha_fin:
        return jsonify({'error': 'Se requieren fechas de inicio y fin'}), 400
    try:
        fecha_inicio = datetime.strptime(fecha_inicio, '%Y-%m-%d').date()
        fecha_fin = datetime.strptime(fecha_fin, '%Y-%m-%d').date()
    except ValueError:
        return jsonify({'error': 'Formato de fecha inválido'}), 400
    
    try:
        query = db.session.query(
            Compras1.IdCentroCosto.label('CentroDeCostos'),
            Compras1.Nit,
            Proveedores.RazonSocial,
            Compras1.Numero.label('Compra'),
            Compras1.NumFactura.label('FraProveedor'),
            Compras1.Fecha,
            Compras2.IdReferencia.label('Codigo'),
            Compras2.Descripcion.label('Producto'),
            Compras2.Cantidad,
            Compras2.Valor,
            (Compras2.Valor / (1 + Compras2.IVA / 100)).label('ValorSinIVA'),
            Compras2.IVA.label('PorcentajeIVA'),
            Compras2.Descuento.label('PorcentajeDescuento'),
            (Compras2.Cantidad * Compras2.Valor).label('Subtotal'),
            (Compras2.Cantidad * Compras2.Valor * (1 - Compras2.Descuento / 100)).label('Total')
        ).join(Compras2, Compras1.Numero == Compras2.Numero
        ).outerjoin(Proveedores, Compras1.Nit == Proveedores.Nit
        ).filter(Compras1.Fecha.between(fecha_inicio, fecha_fin)
        ).order_by(Compras1.Fecha, Compras1.Numero)

        resultado = query.all()
        
        compras = [{
            'CentroDeCostos': r.CentroDeCostos,
            'Nit': r.Nit,
            'RazonSocial': r.RazonSocial,
            'Compra': r.Compra,
            'FraProveedor': r.FraProveedor,
            'Fecha': r.Fecha.strftime('%Y-%m-%d') if r.Fecha else None,
            'Codigo': r.Codigo,
            'Producto': r.Producto,
            'Cantidad': float(r.Cantidad) if r.Cantidad is not None else None,
            'Valor': float(r.Valor) if r.Valor is not None else None,
            'ValorSinIVA': float(r.ValorSinIVA) if r.ValorSinIVA is not None else None,
            'PorcentajeIVA': float(r.PorcentajeIVA) if r.PorcentajeIVA is not None else None,
            'PorcentajeDescuento': float(r.PorcentajeDescuento) if r.PorcentajeDescuento is not None else None,
            'Subtotal': float(r.Subtotal) if r.Subtotal is not None else None,
            'Total': float(r.Total) if r.Total is not None else None
        } for r in resultado]

        app.logger.info(f'Número de compras encontradas: {len(compras)}')
       
        return jsonify({
            'success': True,
            'compras': compras,
            'total_compras': len(compras)
        })
    except Exception as e:
        app.logger.error(f'Error en get_compras_costos_informes: {str(e)}')
        return jsonify({'success': False, 'error': f'Error al obtener los datos: {str(e)}'}), 500

def obtener_datos_compras_costos(fecha_inicio, fecha_fin):
    # Convertir las fechas de string a objetos datetime
    fecha_inicio = datetime.strptime(fecha_inicio, '%Y-%m-%d').date()
    fecha_fin = datetime.strptime(fecha_fin, '%Y-%m-%d').date()
   
    # Consulta para obtener las compras
    compras = db.session.query(
        Compras1.IdCentroCosto.label('CentroDeCostos'),
        Compras1.Nit,
        Proveedores.RazonSocial,
        Compras1.Numero.label('Compra'),
        Compras1.NumFactura.label('FraProveedor'),
        Compras1.Fecha,
        Compras2.IdReferencia.label('Codigo'),
        Compras2.Descripcion.label('Producto'),
        Compras2.Cantidad,
        Compras2.Valor,
        (Compras2.Valor / (1 + Compras2.IVA / 100)).label('ValorSinIVA'),
        Compras2.IVA.label('PorcentajeIVA'),
        Compras2.Descuento.label('PorcentajeDescuento'),
        (Compras2.Cantidad * Compras2.Valor).label('Subtotal'),
        (Compras2.Cantidad * Compras2.Valor * (1 - Compras2.Descuento / 100)).label('Total')
    ).join(
        Compras2,
        Compras1.Numero == Compras2.Numero
    ).outerjoin(
        Proveedores,
        Compras1.Nit == Proveedores.Nit
    ).filter(
        Compras1.Fecha.between(fecha_inicio, fecha_fin)
    ).order_by(
        Compras1.Fecha,
        Compras1.Numero
    ).all()
   
    # Convertir los resultados a un formato serializable
    return [{
        'CentroDeCostos': c.CentroDeCostos or '',
        'Nit': c.Nit,
        'RazonSocial': c.RazonSocial or '',
        'Compra': c.Compra,
        'FraProveedor': c.FraProveedor,
        'Fecha': c.Fecha.strftime('%Y-%m-%d') if c.Fecha else '',
        'Codigo': c.Codigo,
        'Producto': c.Producto,
        'Cantidad': float(c.Cantidad),
        'Valor': float(c.Valor),
        'ValorSinIVA': float(c.ValorSinIVA),
        'PorcentajeIVA': float(c.PorcentajeIVA),
        'PorcentajeDescuento': float(c.PorcentajeDescuento),
        'Subtotal': float(c.Subtotal),
        'Total': float(c.Total)
    } for c in compras]

@app.route('/api/exportar_compras_costos_excel')
@cross_origin(supports_credentials=True)
def exportar_compras_costos_excel():
    try:
        fecha_inicio = request.args.get('fecha_inicio')
        fecha_fin = request.args.get('fecha_fin')
        
        current_app.logger.info(f"Iniciando exportación de Compras x Costos: {fecha_inicio} - {fecha_fin}")
        
        if not fecha_inicio or not fecha_fin:
            raise ValueError("Fechas de inicio y fin son requeridas")

        compras = obtener_datos_compras_costos(fecha_inicio, fecha_fin)
        
        current_app.logger.info(f"Datos obtenidos: {len(compras)} compras")
        
        if not compras:
            return jsonify({"message": "No hay datos para exportar en el rango de fechas especificado"}), 404

        wb = Workbook()
        ws = wb.active
        ws.title = "Compras x Costos"

        # Estilos
        title_font = Font(name='Arial', size=16, bold=True)
        header_font = Font(name='Arial', size=12, bold=True)
        
        # Título
        ws.merge_cells('A1:O1')
        ws['A1'] = "Informe de Compras x Costos"
        ws['A1'].font = title_font
        ws['A1'].alignment = Alignment(horizontal='center')

        # Información de la empresa
        ws['A3'] = "CCD INGENIERIA Y CONSTRUCCIONES S.A.S."
        ws['A4'] = "901092189-5"
        ws['A5'] = "CCD INGENIERIA Y CONSTRUCCIONES"
        
        ws['M3'] = f"Fecha: {fecha_inicio} - {fecha_fin}"
        ws['M4'] = "Dirección: CR 78 A 45 G G 14"
        ws['M5'] = "Teléfono: 3043821361"

        # Encabezados
        headers = ['Centro de Costos', 'NIT', 'Razón Social', 'Compra', 'Fra. Proveedor', 'Fecha', 'Código', 'Producto', 
                   'Cantidad', 'Valor', 'Valor Sin IVA', '% IVA', '% Descuento', 'Subtotal', 'Total']
        for col, header in enumerate(headers, start=1):
            cell = ws.cell(row=7, column=col, value=header)
            cell.font = header_font
            cell.fill = PatternFill(start_color="4F81BD", end_color="4F81BD", fill_type="solid")
            cell.alignment = Alignment(horizontal='center')

        # Datos
        for row, compra in enumerate(compras, start=8):
            ws.cell(row=row, column=1, value=compra['CentroDeCostos'])
            ws.cell(row=row, column=2, value=compra['Nit'])
            ws.cell(row=row, column=3, value=compra['RazonSocial'])
            ws.cell(row=row, column=4, value=compra['Compra'])
            ws.cell(row=row, column=5, value=compra['FraProveedor'])
            ws.cell(row=row, column=6, value=compra['Fecha'])
            ws.cell(row=row, column=7, value=compra['Codigo'])
            ws.cell(row=row, column=8, value=compra['Producto'])
            ws.cell(row=row, column=9, value=compra['Cantidad'])
            ws.cell(row=row, column=10, value=compra['Valor'])
            ws.cell(row=row, column=11, value=compra['ValorSinIVA'])
            ws.cell(row=row, column=12, value=compra['PorcentajeIVA'])
            ws.cell(row=row, column=13, value=compra['PorcentajeDescuento'])
            ws.cell(row=row, column=14, value=compra['Subtotal'])
            ws.cell(row=row, column=15, value=compra['Total'])

        # Ajustar el ancho de las columnas
        for column in ws.columns:
            max_length = 0
            column_letter = get_column_letter(column[0].column)
            for cell in column:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            adjusted_width = (max_length + 2)
            ws.column_dimensions[column_letter].width = adjusted_width

        current_app.logger.info("Archivo Excel creado exitosamente")
        
        output = io.BytesIO()
        wb.save(output)
        output.seek(0)

        response = make_response(output.getvalue())
        response.headers['Content-Disposition'] = 'attachment; filename=Compras_x_Costos.xlsx'
        response.headers['Content-type'] = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        response.headers['Access-Control-Allow-Origin'] = '*'
        
        current_app.logger.info("Archivo Excel enviado exitosamente")
        return response

    except ValueError as ve:
        current_app.logger.error(f"Error de validación: {str(ve)}")
        return jsonify({"error": str(ve)}), 400
    except Exception as e:
        current_app.logger.error(f"Error al exportar Compras x Costos: {str(e)}")
        current_app.logger.error(traceback.format_exc())
        return jsonify({"error": "Error interno del servidor", "details": str(e), "traceback": traceback.format_exc()}), 500

@app.route('/api/informe_compras', methods=['GET'])
@cross_origin()
def get_informe_compras():
    app.logger.info('Función get_informe_compras llamada')
    fecha_inicio = request.args.get('fecha_inicio')
    fecha_fin = request.args.get('fecha_fin')
    tipo_informe = request.args.get('tipo')
   
    app.logger.info(f'Fechas recibidas: inicio={fecha_inicio}, fin={fecha_fin}, tipo={tipo_informe}')
    if not fecha_inicio or not fecha_fin or not tipo_informe:
        return jsonify({'error': 'Se requieren fechas de inicio y fin, y tipo de informe'}), 400
    try:
        fecha_inicio = datetime.strptime(fecha_inicio, '%Y-%m-%d').date()
        fecha_fin = datetime.strptime(fecha_fin, '%Y-%m-%d').date()
    except ValueError:
        return jsonify({'error': 'Formato de fecha inválido'}), 400
   
    try:
        if tipo_informe == 'detalladas':
            query = db.session.query(
                Compras1.Nit,
                Proveedores.RazonSocial,
                Compras1.Numero,
                Compras1.NumFactura,
                Compras1.Fecha,
                Compras2.IdReferencia,
                Compras2.Descripcion,
                Compras2.Cantidad,
                Compras2.Valor.label('ValorUND'),
                Compras2.Descuento,
                Compras2.IVA,
                (Compras2.Valor / (1 + Compras2.IVA / 100)).label('ValorSinIVA'),
                (Compras2.Cantidad * Compras2.Valor).label('ValorBase'),
                (Compras2.Cantidad * Compras2.Valor * (Compras2.Descuento / 100)).label('ValorDescuento'),
                (Compras2.Cantidad * Compras2.Valor * (1 - Compras2.Descuento / 100)).label('ValorBaseDesc'),
                (Compras2.Cantidad * Compras2.Valor).label('Subtotal'),
                (Compras2.Cantidad * Compras2.Valor * (Compras2.IVA / 100)).label('ValorIVA'),
                Compras2.imp_icui.label('ValorICUI'),
                Compras2.ipc.label('ValorIPC'),
                Compras2.imp_ibua.label('ValorIBUA'),
                (Compras2.Cantidad * Compras2.Valor * (1 + Compras2.IVA / 100)).label('Total'),
                CentroCostos.CentroCosto.label('CCostos')
            ).join(Compras2, Compras1.Numero == Compras2.Numero
            ).outerjoin(Proveedores, Compras1.Nit == Proveedores.Nit
            ).outerjoin(CentroCostos, Compras1.IdCentroCosto == CentroCostos.IdCentroCosto
            ).filter(Compras1.Fecha.between(fecha_inicio, fecha_fin), Compras1.Anulado == False
            ).order_by(Compras1.Numero)
        elif tipo_informe in ['porProveedor', 'porValor']:
            query = db.session.query(
                Compras1.Nit,
                Proveedores.RazonSocial,
                func.sum(Compras1.total).label('Total')  # Cambiado de 'Total' a 'total'
            ).outerjoin(Proveedores, Compras1.Nit == Proveedores.Nit
            ).filter(Compras1.Fecha.between(fecha_inicio, fecha_fin), Compras1.Anulado == False
            ).group_by(Compras1.Nit, Proveedores.RazonSocial
            ).order_by(desc('Total') if tipo_informe == 'porValor' else Proveedores.RazonSocial)
        elif tipo_informe == 'porProducto':
            query = db.session.query(
                Compras2.IdReferencia.label('CodProducto'),
                Referencia.Referencia,
                func.sum(Compras2.Cantidad).label('CantidadComprada'),
                func.sum(Compras2.Cantidad * Compras2.Valor).label('ValorTotal')
            ).join(Compras1, Compras1.Numero == Compras2.Numero
            ).join(Referencia, Compras2.IdReferencia == Referencia.IdReferencia
            ).filter(Compras1.Fecha.between(fecha_inicio, fecha_fin), Compras1.Anulado == False
            ).group_by(Compras2.IdReferencia, Referencia.Referencia)
        else:
            return jsonify({'error': 'Tipo de informe no válido'}), 400

        resultado = query.all()
       
        compras = []
        for r in resultado:
            compra = {}
            for idx, column in enumerate(query.column_descriptions):
                key = column['name']
                value = r[idx]
                if isinstance(value, Decimal):
                    compra[key] = float(value)
                elif isinstance(value, date):
                    compra[key] = value.strftime('%Y-%m-%d')
                else:
                    compra[key] = value
            compras.append(compra)

        app.logger.info(f'Número de registros encontrados: {len(compras)}')
       
        return jsonify({
            'success': True,
            'compras': compras,
            'total_registros': len(compras)
        })
    except Exception as e:
        app.logger.error(f'Error en get_informe_compras: {str(e)}')
        return jsonify({'success': False, 'error': f'Error al obtener los datos: {str(e)}'}), 500

@app.route('/api/grupos', methods=['GET'])
def obtener_grupos():
    grupos = Grupo.query.all()
    return jsonify([{
        'codigo': grupo.IdGrupo,
        'descripcion': grupo.Grupo,
        'estado': grupo.Estado,
        'menupos': grupo.menupos,
        'ultimoCodigo': grupo.ultimoCodigo  # Asumiendo que agregaste este campo a tu modelo Grupo
    } for grupo in grupos])

@app.route('/api/departamentos', methods=['GET'])
def obtener_departamentos():
    departamentos = Departamentos.query.all()
    return jsonify([{
        'IdDepartamento': d.IdDepartamento,
        'Departamento': d.Departamento,
        'IdPais': d.IdPais,
        'Estado': d.Estado
    } for d in departamentos])

@app.route('/api/departamentos', methods=['POST'])
def crear_departamento():
    data = request.json
    nuevo_departamento = Departamentos(
        IdDepartamento=data['IdDepartamento'],
        Departamento=data['Departamento'],
        IdPais=data['IdPais'],
        Estado=data['Estado']
    )
    db.session.add(nuevo_departamento)
    db.session.commit()
    return jsonify({'message': 'Departamento creado exitosamente'}), 201

@app.route('/api/departamentos/<string:id>', methods=['PUT'])
def actualizar_departamento(id):
    departamento = Departamentos.query.get(id)
    if not departamento:
        return jsonify({'message': 'Departamento no encontrado'}), 404
    data = request.json
    departamento.Departamento = data['Departamento']
    departamento.IdPais = data['IdPais']
    departamento.Estado = data['Estado']
    db.session.commit()
    return jsonify({'message': 'Departamento actualizado exitosamente'})

@app.route('/api/departamentos/<string:id>', methods=['DELETE'])
def eliminar_departamento(id):
    departamento = Departamentos.query.get(id)
    if not departamento:
        return jsonify({'message': 'Departamento no encontrado'}), 404
    db.session.delete(departamento)
    db.session.commit()
    return jsonify({'message': 'Departamento eliminado exitosamente'})

@app.route('/api/grupos_subcategorias', methods=['GET'])
def obtener_grupos_subcategorias():
    try:
        grupos = Grupo.query.filter_by(Estado=True).all()
        grupos_data = [{
            'IdGrupo': grupo.IdGrupo,
            'Grupo': grupo.Grupo
        } for grupo in grupos]
        print("Grupos para subcategorías obtenidos:", grupos_data)  # Log para debugging
        return jsonify(grupos_data)
    except Exception as e:
        print(f"Error al obtener grupos para subcategorías: {str(e)}")
        return jsonify({'error': 'Error al obtener grupos para subcategorías'}), 500

@app.route('/api/guardar_entrada', methods=['POST'])
def guardar_entrada():
    try:
        data = request.json
        entrada1_data = data.get('entrada1', {})
        entradas2_data = data.get('entradas2', [])

        logging.info(f"Datos recibidos: {data}")

        # Asegurarse de que el usuario 'MIG' existe
        get_or_create_mig_user()

        # Crear instancia de Entradas1
        nueva_entrada1 = Entradas1(
            Numero=entrada1_data.get('Numero'),
            Mes=entrada1_data.get('Mes'),
            Anulado=entrada1_data.get('Anulado', False),
            IdBodega=entrada1_data.get('IdBodega'),
            Observaciones=entrada1_data.get('Observaciones'),
            FechaCreacion=parser.parse(entrada1_data.get('FechaCreacion')),
            IdUsuario='MIG',
            IdConsecutivo=entrada1_data.get('IdConsecutivo'),
            fecha=parser.parse(entrada1_data.get('fecha')),
            subtotal=Decimal(str(entrada1_data.get('subtotal', '0'))),
            total_iva=Decimal(str(entrada1_data.get('total_iva', '0'))),
            total_impoconsumo=Decimal(str(entrada1_data.get('total_impoconsumo', '0'))),
            total_ipc=Decimal(str(entrada1_data.get('total_ipc', '0'))),
            total_ibua=Decimal(str(entrada1_data.get('total_ibua', '0'))),
            total_icui=Decimal(str(entrada1_data.get('total_icui', '0'))),
            total=Decimal(str(entrada1_data.get('total', '0')))
        )
        db.session.add(nueva_entrada1)

        # Guardar en Entradas2 y actualizar SaldosBodega
        mes_actual = datetime.now().strftime('%Y%m')
        for entrada2 in entradas2_data:
            if not entrada2.get('IdReferencia'):
                logging.warning(f"Entrada ignorada debido a IdReferencia vacío: {entrada2}")
                continue

            nueva_entrada2 = Entradas2(
                ID=entrada2.get('ID'),
                Numero=entrada2.get('Numero'),
                IdReferencia=entrada2.get('IdReferencia'),
                Descripcion=entrada2.get('Descripcion'),
                Cantidad=Decimal(str(entrada2.get('Cantidad', '0'))),
                Valor=Decimal(str(entrada2.get('Valor', '0'))),
                IVA=Decimal(str(entrada2.get('IVA', '0'))),
                Descuento=Decimal('0'),
                idunidad=entrada2.get('idunidad'),
                impoconsumo=Decimal(str(entrada2.get('impoconsumo', '0'))),
                ipc=Decimal(str(entrada2.get('ipc', '0'))),
                imp_ibua=Decimal(str(entrada2.get('imp_ibua', '0'))),
                imp_icui=Decimal(str(entrada2.get('imp_icui', '0')))
            )
            db.session.add(nueva_entrada2)

            # Actualizar SaldosBodega
            saldo = SaldosBodega.query.filter_by(
                IdBodega=entrada1_data['IdBodega'],
                Mes=mes_actual,
                IdReferencia=entrada2['IdReferencia']
            ).first()

            if saldo:
                saldo.Entradas += Decimal(str(entrada2['Cantidad']))
                saldo.Saldo += Decimal(str(entrada2['Cantidad']))
                if entrada2.get('lote'):
                    saldo.lote = entrada2['lote']
            else:
                nuevo_saldo = SaldosBodega(
                    IdBodega=entrada1_data['IdBodega'],
                    Mes=mes_actual,
                    IdReferencia=entrada2['IdReferencia'],
                    Entradas=Decimal(str(entrada2['Cantidad'])),
                    Saldo=Decimal(str(entrada2['Cantidad'])),
                    lote=entrada2.get('lote', '')
                )
                db.session.add(nuevo_saldo)

        db.session.commit()

        # Generar documento Word en memoria
        doc_content = generar_documento_entrada(entrada1_data, entradas2_data)
        filename = f"Entrada_Bodega_{entrada1_data['Numero']}.docx"

        return jsonify({
            'success': True, 
            'message': 'Entrada guardada y saldo actualizado con éxito',
            'documento': filename,
            'doc_content': base64.b64encode(doc_content).decode('utf-8')
        })

    except SQLAlchemyError as e:
        db.session.rollback()
        error_trace = traceback.format_exc()
        logging.error(f"Error de base de datos al guardar la entrada: {str(e)}\n{error_trace}")
        return jsonify({
            'success': False, 
            'message': 'Error de base de datos al guardar la entrada',
            'error': str(e),
            'trace': error_trace
        }), 500
    except Exception as e:
        db.session.rollback()
        error_trace = traceback.format_exc()
        logging.error(f"Error al guardar la entrada: {str(e)}\n{error_trace}")
        return jsonify({
            'success': False, 
            'message': 'Error interno del servidor al guardar la entrada',
            'error': str(e),
            'trace': error_trace
        }), 500

@app.errorhandler(InternalServerError)
def handle_500(e):
    original = getattr(e, "original_exception", None)
    
    if original is None:
        # direct 500 error, such as abort(500)
        return jsonify(error=str(e)), 500

    # wrapped unhandled error
    return jsonify(error=str(original)), 500

@app.route('/api/descargar_documento/<filename>', methods=['GET'])
def descargar_documento(filename):
    try:
        file_path = os.path.join(current_app.config['UPLOAD_FOLDER'], filename)
        if not os.path.exists(file_path):
            current_app.logger.error(f"Archivo no encontrado: {file_path}")
            abort(404, description="Documento no encontrado")

        return send_file(
            file_path,
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except Exception as e:
        current_app.logger.error(f"Error al descargar el documento: {str(e)}")
        abort(500, description=f"Error interno del servidor al descargar el documento: {str(e)}")

def generar_documento_entrada(entrada, productos):
    doc = Document()
    
    # Configurar márgenes de página
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

    # Crear tabla para el encabezado
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.allow_autofit = False
    
    # Ajustar ancho de las columnas
    table.columns[0].width = Inches(2.5)
    table.columns[1].width = Inches(4.5)

    # Celda para el logo
    logo_cell = table.cell(0, 0)
    logo_path = os.path.join(current_app.root_path, 'static', 'img', 'logoEmpresa.png')
    if os.path.exists(logo_path):
        logo_paragraph = logo_cell.paragraphs[0]
        logo_run = logo_paragraph.add_run()
        logo_run.add_picture(logo_path, width=Inches(2))
    else:
        print(f"Logo no encontrado en: {logo_path}")
    
    company_info = logo_cell.add_paragraph()
    company_info.add_run('CCD INGENIERÍA Y CONSTRUCCIONES S.A.S.\n').bold = True
    company_info.add_run('NIT. 901.092.189-5')

    # Celda para el título y la fecha
    title_cell = table.cell(0, 1)
    title_paragraph = title_cell.paragraphs[0]
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    title_run = title_paragraph.add_run('Entrada a Bodega\n')
    title_run.bold = True
    title_run.font.size = Pt(16)
    date_run = title_paragraph.add_run(f"Fecha: {entrada.get('fecha', datetime.now().strftime('%Y-%m-%d'))}")
    date_run.font.size = Pt(10)

    # Agregar línea separadora
    doc.add_paragraph('_' * 100)

    # Resto del documento
    doc.add_heading('Observaciones', level=1)
    doc.add_paragraph(entrada.get('Observaciones', 'Sin observaciones'))

    doc.add_heading('Productos', level=1)
    products_table = doc.add_table(rows=1, cols=5)
    products_table.style = 'Table Grid'
    hdr_cells = products_table.rows[0].cells
    hdr_cells[0].text = 'ID Referencia'
    hdr_cells[1].text = 'Descripción'
    hdr_cells[2].text = 'Cantidad'
    hdr_cells[3].text = 'Valor'
    hdr_cells[4].text = 'Subtotal'

    for producto in productos:
        row_cells = products_table.add_row().cells
        row_cells[0].text = producto.get('IdReferencia', '')
        row_cells[1].text = producto.get('Descripcion', '')
        row_cells[2].text = str(producto.get('Cantidad', 0))
        row_cells[3].text = str(producto.get('Valor', 0))
        subtotal = float(producto.get('Cantidad', 0)) * float(producto.get('Valor', 0))
        row_cells[4].text = str(subtotal)

    doc.add_heading('Firma', level=1)
    doc.add_paragraph('_' * 30)
    doc.add_paragraph('Nombre y Firma')

    # Guardar en un objeto BytesIO
    from io import BytesIO
    doc_buffer = BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)
    return doc_buffer.getvalue()

# Asegúrate de que este directorio exista y tenga los permisos adecuados
if not os.path.exists(UPLOAD_FOLDER):
    try:
        os.makedirs(UPLOAD_FOLDER)
    except Exception as e:
        print(f"Error al crear el directorio UPLOAD_FOLDER: {e}")
        raise

# Verifica los permisos
if not os.access(UPLOAD_FOLDER, os.W_OK):
    print(f"No se tienen permisos de escritura en {UPLOAD_FOLDER}")
    raise PermissionError(f"No se tienen permisos de escritura en {UPLOAD_FOLDER}")

print(f"UPLOAD_FOLDER configurado en: {UPLOAD_FOLDER}")


@app.route('/api/importar-productos', methods=['POST'])
def importar_productos():
    productos = request.json
    for producto in productos:
        nueva_referencia = Referencia(
            IdReferencia=producto['IdReferencia'],
            Referencia=producto['Referencia'],
            Marca=producto['Marca'],
            EstadoProducto=producto['EstadoProducto'],
            IdGrupo=producto['IdGrupo'],
            IdUnidad=producto['IdUnidad'],
            Ubicacion=producto['Ubicacion'],
            Costo=producto['Costo'],
            PrecioVenta1=producto['PrecioVenta1'],
            IVA=producto['IVA'],
            Estado=producto['Estado'],
            Tipo=producto['Tipo'],
            ManejaInventario=producto['ManejaInventario']
        )
        db.session.add(nueva_referencia)
    
    try:
        db.session.commit()
        return jsonify({'message': 'Productos importados exitosamente'}), 200
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': f'Error al importar productos: {str(e)}'}), 500

@app.route('/api/subcategorias', methods=['GET', 'POST'])
def manejar_subcategorias():
    if request.method == 'GET':
        subcategorias = Subcategorias.query.all()
        return jsonify([{
            'idsubcategoria': s.idsubcategoria,
            'categoria': s.categoria,
            'idgrupo': s.idgrupo,
            'idsubgrupo': s.idsubgrupo,
            'estado': s.estado
        } for s in subcategorias])
    
    elif request.method == 'POST':
        data = request.json
        nueva_subcategoria = Subcategorias(
            idsubcategoria=data['idsubcategoria'],
            categoria=data['categoria'],
            idgrupo=data['idgrupo'],
            idsubgrupo=data['idsubgrupo'],
            estado=data['estado']
        )
        db.session.add(nueva_subcategoria)
        try:
            db.session.commit()
            return jsonify({'message': 'Subcategoría creada exitosamente'}), 201
        except Exception as e:
            db.session.rollback()
            return jsonify({'error': str(e)}), 400

@app.route('/api/subcategorias/<string:id>', methods=['GET', 'PUT', 'DELETE'])
def manejar_subcategoria_individual(id):
    subcategoria = Subcategorias.query.get_or_404(id)
    
    if request.method == 'GET':
        return jsonify({
            'idsubcategoria': subcategoria.idsubcategoria,
            'categoria': subcategoria.categoria,
            'idgrupo': subcategoria.idgrupo,
            'idsubgrupo': subcategoria.idsubgrupo,
            'estado': subcategoria.estado
        })
    
    elif request.method == 'PUT':
        data = request.json
        subcategoria.categoria = data['categoria']
        subcategoria.idgrupo = data['idgrupo']
        subcategoria.idsubgrupo = data['idsubgrupo']
        subcategoria.estado = data['estado']
        try:
            db.session.commit()
            return jsonify({'message': 'Subcategoría actualizada exitosamente'})
        except Exception as e:
            db.session.rollback()
            return jsonify({'error': str(e)}), 400
    
    elif request.method == 'DELETE':
        db.session.delete(subcategoria)
        try:
            db.session.commit()
            return jsonify({'message': 'Subcategoría eliminada exitosamente'})
        except Exception as e:
            db.session.rollback()
            return jsonify({'error': str(e)}), 400

@app.route('/api/obtener_grupos_subgrupos', methods=['GET'])
def obtener_grupos_subgrupos():
    try:
        grupos = Grupo.query.filter_by(Estado=True).all()
        grupos_data = [{
            'codigo': grupo.IdGrupo,
            'descripcion': grupo.Grupo
        } for grupo in grupos]
        return jsonify(grupos_data)
    except Exception as e:
        print(f"Error al obtener grupos para subgrupos: {str(e)}")
        return jsonify({'error': 'Error al obtener grupos para subgrupos'}), 500

@app.route('/api/subgrupos', methods=['GET'])
def obtener_subgrupos():
    try:
        subgrupos = SubGrupos.query.all()
        subgrupos_data = [{
            'IdSubgrupo': subgrupo.IdSubgrupo,
            'Subgrupo': subgrupo.Subgrupo,
            'IdGrupo': subgrupo.IdGrupo,
            'Estado': subgrupo.Estado
        } for subgrupo in subgrupos]
        return jsonify(subgrupos_data)
    except Exception as e:
        print(f"Error al obtener subgrupos: {str(e)}")
        return jsonify({'error': 'Error al obtener subgrupos'}), 500

@app.route('/api/subgrupos/<string:id_grupo>', methods=['GET'])
def obtener_subgrupos_por_grupo(id_grupo):
    try:
        subgrupos = SubGrupos.query.filter_by(IdGrupo=id_grupo, Estado=True).all()
        subgrupos_data = [{
            'IdSubgrupo': subgrupo.IdSubgrupo,
            'Subgrupo': subgrupo.Subgrupo
        } for subgrupo in subgrupos]
        return jsonify(subgrupos_data)
    except Exception as e:
        print(f"Error al obtener subgrupos: {str(e)}")
        return jsonify({'error': 'Error al obtener subgrupos'}), 500

@app.route('/api/grupos/<string:id>', methods=['DELETE'])
def eliminar_grupo(id):
    try:
        grupo = Grupo.query.get(id)
        if grupo:
            db.session.delete(grupo)
            db.session.commit()
            return jsonify({'success': True, 'message': 'Grupo eliminado exitosamente'})
        else:
            return jsonify({'success': False, 'message': 'Grupo no encontrado'}), 404
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'message': str(e)}), 500

@app.route('/api/grupos/<codigo>', methods=['GET'])
def verificar_grupo(codigo):
    grupo = Grupo.query.get(codigo)
    return jsonify({'exists': grupo is not None})

def verificar_referencias_existen(referencias):
    existentes = Referencia.query.filter(Referencia.IdReferencia.in_(referencias)).with_entities(Referencia.IdReferencia).all()
    existentes = set(r[0] for r in existentes)
    return all(ref in existentes for ref in referencias)

@app.route('/maestros/bodegas')
def pagina_bodegas():
    return render_template('bodegas.html')

@app.route('/api/configuracion', methods=['GET', 'POST'])
def manejar_configuracion():
    if request.method == 'POST':
        data = request.json
        # Procesar data...
        return jsonify({'message': 'Configuración actualizada exitosamente'}), 200
    else:
        configuracion = {
            'empresa': 'MIG-ALMACEN, TEXTIL',
            'licencia': 'XXXX-XXXX-XXXX-XXXX',
            # ... más configuraciones ...
        }
        return jsonify(configuracion)

def generar_licencia(caracteristicas_equipo, nit):
    caracteristicas = base64.b64decode(caracteristicas_equipo).decode('utf-8')
    clave_secreta = "MIGSistemas2024"
    datos = caracteristicas + nit + clave_secreta
    hash_objeto = hashlib.sha256(datos.encode())
    hash_bytes = hash_objeto.digest()
    licencia_base = base64.b64encode(hash_bytes).decode()[:16]
    caracteres_aleatorios = ''.join(random.choices(string.ascii_uppercase + string.digits, k=8))
    licencia = f"{licencia_base}-{caracteres_aleatorios}-{nit[-4:]}"
    return licencia

def enviar_correo(licencia, usuario, password):
    try:
        msg = Message('Nueva Licencia Generada',
                    sender=app.config['MAIL_DEFAULT_SENDER'],
                    recipients=['riosjuan3053@gmail.com', 'johnsod8729@gmail.com'])
        msg.body = f"""
        Se ha generado una nueva licencia:
        
        NIT: {licencia.nit}
        Razón Social: {licencia.razonsocial}
        Nombre Comercial: {licencia.nombrecomercial}
        Número de Licencia: {licencia.numerolicencia}
        Tipo de Licencia: {licencia.tipolicencia}
        Fecha de Vencimiento: {licencia.fechavencimiento if licencia.fechavencimiento else 'N/A'}

        Información de acceso:
        Usuario: {usuario}
        Contraseña: {password}

        Por favor, cambie su contraseña después del primer inicio de sesión.
        """
        logging.info(f"Intentando enviar correo a {', '.join(msg.recipients)}")
        mail.send(msg)
        logging.info("Correo enviado exitosamente")
        return True
    except Exception as e:
        logging.error(f"Error al enviar correo: {str(e)}")
        logging.error(f"Detalles del error: {traceback.format_exc()}")
        return False
    
@app.before_request
def log_request_info():
    app.logger.info('Ruta solicitada: %s %s', request.method, request.path)

@app.after_request
def log_response_info(response):
    app.logger.debug('Response Status: %s', response.status)
    app.logger.debug('Response: %s', response.get_data())
    return response

@app.after_request
def after_request(response):
    response.headers.add('Access-Control-Allow-Origin', '*')
    response.headers.add('Access-Control-Allow-Headers', 'Content-Type,Authorization')
    response.headers.add('Access-Control-Allow-Methods', 'GET,PUT,POST,DELETE,OPTIONS')
    return response

@app.route('/api/solicitar_licencia', methods=['POST'])
def solicitar_licencia():
    try:
        data = request.json
        logging.info(f"Datos recibidos: {data}")
        
        # Validar datos recibidos
        campos_requeridos = ['nit', 'razonsocial', 'nombrecomercial', 'ubicacioncomercial', 'ciudad', 'telefono', 'version', 'cantidadusuario', 'tipolicencia', 'caracteristicas_equipo']
        for campo in campos_requeridos:
            if campo not in data or not data[campo]:
                return jsonify({'success': False, 'message': f'El campo {campo} es requerido y no puede estar vacío'}), 400

        # Generar licencia
        numero_licencia = generar_licencia(data['caracteristicas_equipo'], data['nit'])
        logging.info(f"Número de licencia generado: {numero_licencia}")
        
        # Crear nueva licencia en la base de datos
        nueva_licencia = Licencia(
            id=numero_licencia,  # Usar el número de licencia como ID
            nit=data['nit'],
            razonsocial=data['razonsocial'],
            nombrecomercial=data['nombrecomercial'],
            ubicacioncomercial=data['ubicacioncomercial'],
            ciudad=data['ciudad'],
            telefono=data['telefono'],
            version=data['version'],
            numerolicencia=numero_licencia,
            cantidadusuario=data['cantidadusuario'],
            tipolicencia=data['tipolicencia'],
            fechavencimiento=datetime.strptime(data['fechavencimiento'], '%Y-%m-%d').date() if data['tipolicencia'] == 'RENTA' and 'fechavencimiento' in data else None
        )
        db.session.add(nueva_licencia)

        # Crear usuario administrador
        usuario = data['nit']  # Usar NIT como nombre de usuario
        password = generar_password()
        nuevo_usuario = Usuarios(
            IdUsuario=usuario,
            Contraseña=password,
            Descripcion=f"Usuario administrador para {data['razonsocial']}",
            Estado=True,
            NivelAcceso=100  # Nivel máximo de acceso
        )
        db.session.add(nuevo_usuario)

        # Asignar todos los permisos disponibles al usuario
        todos_los_permisos = Permisos.query.all()
        for permiso in todos_los_permisos:
            nuevo_permiso_usuario = UsuariosPermisos(IdUsuario=usuario, IdPermiso=permiso.IdPermiso)
            db.session.add(nuevo_permiso_usuario)

        db.session.commit()
        logging.info("Licencia, usuario administrador y permisos creados y guardados en la base de datos")

        # Enviar correo
        correo_enviado = enviar_correo(nueva_licencia, usuario, password)
        if not correo_enviado:
            logging.warning("El correo no pudo ser enviado, pero la licencia y el usuario fueron creados.")

        return jsonify({
            'success': True,
            'message': 'Licencia creada exitosamente',
            'correoEnviado': correo_enviado,
            'licencia': {
                'nit': nueva_licencia.nit,
                'razonsocial': nueva_licencia.razonsocial,
                'nombrecomercial': nueva_licencia.nombrecomercial,
                'numerolicencia': nueva_licencia.numerolicencia,
                'tipolicencia': nueva_licencia.tipolicencia,
                'fechavencimiento': nueva_licencia.fechavencimiento.isoformat() if nueva_licencia.fechavencimiento else None
            },
            'usuario': {
                'IdUsuario': usuario,
                'password': password  # Nota: En producción, no deberías devolver la contraseña
            }
        }), 201
    except SQLAlchemyError as e:
        db.session.rollback()
        logging.error(f"Error de base de datos: {str(e)}")
        return jsonify({'success': False, 'message': f'Error de base de datos: {str(e)}'}), 500
    except Exception as e:
        db.session.rollback()
        logging.error(f"Error al procesar la solicitud: {str(e)}")
        logging.error(traceback.format_exc())
        return jsonify({'success': False, 'message': f'Error al procesar la solicitud: {str(e)}'}), 500

@app.route('/api/verificar_licencia', methods=['POST'])
def verificar_licencia():
    data = request.json
    licencia_ingresada = data['licencia']
    nit = data['nit']
    
    # Buscar la licencia en la base de datos
    licencia = Licencia.query.filter_by(id=licencia_ingresada, nit=nit).first()
    
    if licencia:
        # Aquí puedes agregar más verificaciones si es necesario
        # Por ejemplo, verificar la fecha de vencimiento para licencias de RENTA
        if licencia.tipolicencia == 'RENTA':
            if licencia.fechavencimiento and licencia.fechavencimiento < datetime.now().date():
                return jsonify({'valida': False, 'message': 'La licencia ha expirado'})
        
        return jsonify({'valida': True})
    else:
        return jsonify({'valida': False, 'message': 'Licencia no encontrada'})

@app.route('/api/entradas_inventario', methods=['POST'])
def crear_entrada_inventario():
    data = request.json
    
    nueva_entrada = Entradas1(
        Numero=data['Numero'],
        Mes=data['Mes'],
        IdBodega=data['IdBodega'],
        Observaciones=data['Observaciones'],
        IdUsuario=data['IdUsuario'],
        IdConsecutivo=data['IdConsecutivo']
    )
    db.session.add(nueva_entrada)
    
    for detalle in data['detalles']:
        nuevo_detalle = Entradas2(
            ID=f"{data['Numero']}_{detalle['IdReferencia']}",
            Numero=data['Numero'],
            IdReferencia=detalle['IdReferencia'],
            Descripcion=detalle['Descripcion'],
            Cantidad=detalle['Cantidad'],
            Valor=detalle['Valor'],
            Idunidad=detalle['Idunidad']
        )
        db.session.add(nuevo_detalle)
    
    db.session.commit()
    return jsonify({'message': 'Entrada de inventario creada exitosamente'}), 201

@app.route('/api/consecutivos_disponibles', methods=['GET'])
def obtener_consecutivos_disponibles():
    consecutivos = Consecutivos.query.filter_by(Estado=True).all()
    return jsonify([{'IdConsecutivo': c.IdConsecutivo, 'Consecutivo': c.Consecutivo} for c in consecutivos])

@app.route('/api/consecutivos', methods=['GET'])
def obtener_consecutivos():
    consecutivos = Consecutivos.query.all()
    return jsonify([{
        'IdConsecutivo': c.IdConsecutivo,
        'Consecutivo': c.Consecutivo,
        'Formulario': c.Formulario,
        'Prefijo': c.Prefijo,
        'Desde': c.Desde,
        'Hasta': c.Hasta,
        'Actual': c.Actual,
        'Resolucion': c.Resolucion,
        'FechaResolucion': c.FechaResolucion.strftime('%Y-%m-%d') if c.FechaResolucion else None,
        'ObservacionesResolucion': c.ObservacionesResolucion,
        'Estado': c.Estado,
        'Comprobante': c.Comprobante,
        'fechafinresolucion': c.fechafinresolucion,
        'tiporesolucion': c.tiporesolucion
    } for c in consecutivos])

@app.route('/api/bodegas_disponibles', methods=['GET'])
def obtener_bodegas_disponibles():
    bodegas = Bodegas.query.filter_by(Estado=True).all()
    return jsonify([{'IdBodega': b.IdBodega, 'Descripcion': b.Descripcion} for b in bodegas])

@app.route('/api/consecutivos_entradas_inventario', methods=['GET'])
def obtener_consecutivos_entradas_inventario():
    try:
        # Suponiendo que 'EI' es el código para Entradas de Inventario
        consecutivos = Consecutivos.query.filter_by(Formulario='EI', Estado=True).all()
        return jsonify([{
            'IdConsecutivo': c.IdConsecutivo,
            'Descripcion': c.Consecutivo,
            'Prefijo': c.Prefijo,
            'Actual': c.Actual
        } for c in consecutivos])
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/ultimo_consecutivo_entradas', methods=['GET'])
def ultimo_consecutivo_entradas():
    try:
        consecutivo = Consecutivos.query.filter_by(Formulario='EI').first()
        if consecutivo:
            ultimo_consecutivo = f"{consecutivo.Prefijo}{consecutivo.Actual.zfill(2)}"
            return jsonify({'success': True, 'ultimoConsecutivo': ultimo_consecutivo})
        else:
            return jsonify({'success': False, 'message': 'No se encontró el consecutivo para Entradas de Inventario'})
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)})

@app.route('/api/salidas_inventario_informes', methods=['GET'])
@cross_origin()
def get_salidas_inventario_informes():
    app.logger.info('Función get_salidas_inventario_informes llamada')
    fecha_inicio = request.args.get('fecha_inicio')
    fecha_fin = request.args.get('fecha_fin')
    
    app.logger.info(f'Fechas recibidas: inicio={fecha_inicio}, fin={fecha_fin}')

    if not fecha_inicio or not fecha_fin:
        return jsonify({'error': 'Se requieren fechas de inicio y fin'}), 400

    try:
        fecha_inicio = datetime.strptime(fecha_inicio, '%Y-%m-%d').date()
        fecha_fin = datetime.strptime(fecha_fin, '%Y-%m-%d').date()
    except ValueError:
        return jsonify({'error': 'Formato de fecha inválido'}), 400

    try:
        # Primero, obtenemos los datos de Salidas1
        salidas1 = Salidas1.query.filter(Salidas1.fecha.between(fecha_inicio, fecha_fin)).all()
        
        resultado = []
        for s1 in salidas1:
            # Luego, para cada salida en Salidas1, buscamos las correspondientes en Salidas2
            salidas2 = Salidas2.query.filter_by(Numero=s1.Numero).all()
            for s2 in salidas2:
                resultado.append({
                    'Numero': s1.Numero,
                    'Fecha': s1.FechaCreacion.strftime('%Y-%m-%d %H:%M:%S') if s1.FechaCreacion else None,
                    'IdReferencia': s2.IdReferencia,
                    'Descripcion': s2.Descripcion,
                    'Cantidad': float(s2.Cantidad) if s2.Cantidad is not None else None,
                    'Valor': float(s2.Valor) if s2.Valor is not None else None,
                    'Total': float(s1.total) if s1.total is not None else None
                })

        app.logger.info(f'Número de salidas encontradas: {len(resultado)}')
        
        return jsonify({
            'salidas': resultado,
            'total_salidas': len(resultado)
        })
    except Exception as e:
        app.logger.error(f'Error en get_salidas_inventario_informes: {str(e)}')
        return jsonify({'error': f'Error al obtener los datos: {str(e)}'}), 500

@app.route('/api/saldos_bodegas', methods=['GET'])
def get_saldos_bodegas():
    id_referencia = request.args.get('idReferencia')
    id_bodega = request.args.get('idBodega')
    mes_actual = datetime.now().strftime('%Y%m')

    saldo = SaldosBodega.query.filter_by(
        IdBodega=id_bodega,
        Mes=mes_actual,
        IdReferencia=id_referencia
    ).first()

    if saldo:
        return jsonify([{
            'IdReferencia': saldo.IdReferencia,
            'IdBodega': saldo.IdBodega,
            'Saldo': float(saldo.Saldo)
        }])
    else:
        return jsonify([])

@app.route('/api/licencia/<nit>', methods=['GET'])
def obtener_licencia_por_nit(nit):
    licencia = Licencia.query.filter_by(nit=nit).first()
    if licencia:
        return jsonify({
            'success': True,
            'licencia': {
                'nit': licencia.nit,
                'razonSocial': licencia.razon_social,
                'nombreComercial': licencia.nombre_comercial,
                'ubicacionComercial': licencia.ubicacion_comercial,
                'ciudad': licencia.ciudad,
                'telefono': licencia.telefono,
                'version': licencia.version,
                'numeroLicencia': licencia.numero_licencia,
                'cantidadUsuarios': licencia.cantidad_usuarios,
                'tipoLicencia': licencia.tipo_licencia,
                'fechaVencimiento': licencia.fecha_vencimiento.isoformat() if licencia.fecha_vencimiento else None
            }
        })
    else:
        return jsonify({'success': False, 'message': 'No se encontró licencia para el NIT proporcionado'}), 404

def generar_numero_orden_compra():
    consecutivo = Consecutivos.query.filter_by(Formulario='OC').first()
    if not consecutivo:
        raise ValueError("No se encontró el consecutivo para Órdenes de Compra")
    
    nuevo_numero = f"{consecutivo.Prefijo}{int(consecutivo.Actual):04d}"
    
    # Incrementar el consecutivo
    consecutivo.Actual = str(int(consecutivo.Actual) + 1).zfill(4)
    db.session.commit()
    
    return nuevo_numero

@app.route('/api/guardar_orden_compra', methods=['POST'])
def guardar_orden_compra():
    try:
        data = request.json
        orden_compra_data = data.get('ordenCompra', {})
        detalles_data = data.get('detalles', [])

        nit_proveedor = orden_compra_data.get('proveedor')
        if not nit_proveedor:
            return jsonify({'success': False, 'message': 'El Nit del proveedor es requerido'}), 400

        proveedor = Proveedores.query.get(nit_proveedor)
        if not proveedor:
            return jsonify({'success': False, 'message': 'Proveedor no encontrado'}), 404

        nuevo_numero = generar_numero_orden_compra()

        nueva_orden = OrdenesCompra1(
            Numero=nuevo_numero,
            Mes=datetime.now().strftime('%Y%m'),
            Nit=nit_proveedor,
            FechaCreacion=datetime.now(),
            IdUsuario=session.get('user_id', 'usuario_default'),
            Observaciones=orden_compra_data.get('observaciones', ''),
            IdBodega=orden_compra_data.get('bodega'),
            IdConsecutivo=orden_compra_data.get('idConsecutivo')
        )

        try:
            nueva_orden.Descuento = Decimal(str(orden_compra_data.get('valorDescuento', '0')))
        except InvalidOperation:
            nueva_orden.Descuento = Decimal('0')

        db.session.add(nueva_orden)

        detalles_guardados = []
        for detalle in detalles_data:
            referencia = Referencia.query.get(detalle['idReferencia'])
            if not referencia:
                return jsonify({
                    'success': False,
                    'message': f"La referencia {detalle['idReferencia']} no existe en la base de datos."
                }), 400
            
            nuevo_detalle = OrdenesCompra2(
                ID=f"{nuevo_numero}_{detalle['idReferencia']}_{uuid.uuid4().hex[:8]}",  # ID único
                Numero=nuevo_numero,
                IdReferencia=detalle['idReferencia'],
                Descripcion=detalle['descripcion'],
                CantidadPedida=Decimal(str(detalle['cantidad'])),
                Valor=Decimal(str(detalle['valorUnitario'])),
                Iva=Decimal(str(detalle.get('iva', '0'))),
                idunidad=detalle.get('unidad', '')
            )
            db.session.add(nuevo_detalle)
            detalles_guardados.append(nuevo_detalle)

        db.session.commit()

        # Enviar correo electrónico
        enviar_correo_orden_compra(nueva_orden, detalles_guardados, proveedor.Email)

        return jsonify({
            'success': True, 
            'message': 'Orden de compra guardada y enviada por correo exitosamente',
            'numero': nuevo_numero
        })

    except Exception as e:
        db.session.rollback()
        print(f"Error al guardar la orden de compra: {str(e)}")
        print(traceback.format_exc())
        return jsonify({
            'success': False, 
            'message': f'Error al guardar la orden de compra: {str(e)}'
        }), 500

def generar_pdf_orden_compra(orden, detalles):
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=0.5*inch, bottomMargin=0.5*inch, leftMargin=0.5*inch, rightMargin=0.5*inch)
    elements = []
    
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name='Center', alignment=1))
    styles.add(ParagraphStyle(name='Right', alignment=2))
    styles.add(ParagraphStyle(name='Left', alignment=0))
    
    # Encabezado
    logo_path = os.path.join(current_app.root_path, 'static', 'img', 'logoEmpresa.png')
    img = Image(logo_path, width=1*inch, height=1*inch)
    header_data = [
        [img, Paragraph("<b>CCD INGENIERÍA Y CONSTRUCCIONES S.A.S.</b><br/>NIT 901.092.189-5<br/>CRA 78A 45 GG 14<br/>Tel: (57) 3023331325 - Ext. undefined<br/>Medellín - Colombia<br/>facturacion.ccding@gmail.com", styles['Left']), Paragraph("<b>Orden de Compra</b><br/>No. " + orden.Numero, styles['Right'])]
    ]
    header_table = Table(header_data, colWidths=[1.2*inch, 4.3*inch, 2*inch])
    header_table.setStyle(TableStyle([
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('ALIGN', (2, 0), (2, 0), 'RIGHT'),
    ]))
    elements.append(header_table)
    elements.append(Spacer(1, 0.2*inch))
    
    # Información del proveedor y fechas (combinadas)
    info_data = [
        ["Señores", orden.proveedor.RazonSocial, "Fecha y hora Orden", ""],
        ["NIT", orden.Nit, "Generación", orden.FechaCreacion.strftime('%d/%m/%Y, %H:%M')],
        ["Dirección", orden.proveedor.Direccion or "", "Expedición", orden.FechaCreacion.strftime('%d/%m/%Y, %H:%M')],
        ["Teléfono", orden.proveedor.Telefono1 or "", "Vencimiento", (orden.FechaCreacion + timedelta(days=30)).strftime('%d/%m/%Y')],
        ["Ciudad", orden.proveedor.IdCiudad or "", "", ""]
    ]
    info_table = Table(info_data, colWidths=[1*inch, 2.7*inch, 1.3*inch, 2.5*inch])
    info_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
        ('BACKGROUND', (0, 1), (1, 1), colors.white),
        ('BACKGROUND', (2, 1), (-1, 1), colors.lightgrey),
        ('BACKGROUND', (0, 2), (1, 2), colors.lightgrey),
        ('BACKGROUND', (2, 2), (-1, 2), colors.white),
        ('BACKGROUND', (0, 3), (1, 3), colors.white),
        ('BACKGROUND', (2, 3), (-1, 3), colors.lightgrey),
        ('BACKGROUND', (0, 4), (1, 4), colors.lightgrey),
        ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 0), (-1, -1), 8),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 1),
        ('TOPPADDING', (0, 0), (-1, -1), 1),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('SPAN', (2, 0), (3, 0))
    ]))
    elements.append(info_table)
    elements.append(Spacer(1, 0.2*inch))
    
    # Tabla de productos
    data = [['Item', 'Descripción', 'Cantidad', 'Vr. Total']]
    for idx, detalle in enumerate(detalles, start=1):
        subtotal = detalle.CantidadPedida * detalle.Valor
        data.append([
            str(idx),
            detalle.Descripcion,
            str(detalle.CantidadPedida),
            f"${subtotal:.2f}"
        ])
    
    table = Table(data, repeatRows=1, colWidths=[0.5*inch, 4*inch, 1*inch, 2*inch])
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 8),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 3),
        ('BACKGROUND', (0, 1), (-1, -1), colors.white),
        ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 1), (-1, -1), 8),
        ('TOPPADDING', (0, 1), (-1, -1), 3),
        ('BOTTOMPADDING', (0, 1), (-1, -1), 3),
        ('GRID', (0, 0), (-1, -1), 1, colors.black)
    ]))
    elements.append(table)
    
    # Totales
    total = sum(detalle.CantidadPedida * detalle.Valor for detalle in detalles)
    elements.append(Spacer(1, 0.1*inch))
    elements.append(Paragraph(f"Total items: {len(detalles)}", styles['Left']))
    totales_data = [
        ["Total Bruto:", f"${total:.2f}"],
        ["Total a Pagar:", f"${total:.2f}"]
    ]
    totales_table = Table(totales_data, colWidths=[5.5*inch, 2*inch])
    totales_table.setStyle(TableStyle([
        ('ALIGN', (0, 0), (0, -1), 'RIGHT'),
        ('ALIGN', (1, 0), (1, -1), 'RIGHT'),
        ('FONTNAME', (0, 0), (-1, -1), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 8),
        ('BACKGROUND', (0, 0), (-1, -1), colors.lightgrey),
        ('GRID', (0, 0), (-1, -1), 1, colors.white)
    ]))
    elements.append(totales_table)
    
    # Valor en letras
    elements.append(Spacer(1, 0.1*inch))
    elements.append(Paragraph(f"Valor en Letras: {numero_a_letras(total)}", styles['Left']))
    
    # Condiciones de pago
    elements.append(Spacer(1, 0.2*inch))
    elements.append(Paragraph("<b>Condiciones de Pago:</b>", styles['Normal']))
    elements.append(Paragraph(f"Consignación    $    {total:.2f}", styles['Normal']))
    
    # Observaciones
    elements.append(Spacer(1, 0.2*inch))
    elements.append(Paragraph("<b>Observaciones:</b>", styles['Normal']))
    elements.append(Paragraph(orden.Observaciones, styles['Normal']))
    
    doc.build(elements)
    pdf = buffer.getvalue()
    buffer.close()
    return pdf

def numero_a_letras(numero):
    unidades = ['', 'UN ', 'DOS ', 'TRES ', 'CUATRO ', 'CINCO ', 'SEIS ', 'SIETE ', 'OCHO ', 'NUEVE ']
    decenas = ['DIEZ ', 'ONCE ', 'DOCE ', 'TRECE ', 'CATORCE ', 'QUINCE ', 'DIECISEIS ', 'DIECISIETE ', 'DIECIOCHO ', 'DIECINUEVE ']
    diez_diez = ['', '', 'VEINTI', 'TREINTA ', 'CUARENTA ', 'CINCUENTA ', 'SESENTA ', 'SETENTA ', 'OCHENTA ', 'NOVENTA ']
    cientos = ['', 'CIENTO ', 'DOSCIENTOS ', 'TRESCIENTOS ', 'CUATROCIENTOS ', 'QUINIENTOS ', 'SEISCIENTOS ', 'SETECIENTOS ', 'OCHOCIENTOS ', 'NOVECIENTOS ']

    def convert_grupo(n):
        if n == '100':
            return 'CIEN '
        
        output = ''
        if len(n) == 3:
            if n[0] != '0':
                output = cientos[int(n[0])]
            n = n[1:]

        if len(n) == 2:
            if n[0] == '1':
                output += decenas[int(n[1])]
            else:
                output += diez_diez[int(n[0])]
                if int(n[1]) > 0:
                    if int(n[0]) > 2:
                        output += 'Y '
                    output += unidades[int(n[1])]
        elif len(n) == 1:
            output += unidades[int(n)]

        return output

    numero = int(round(numero, 2))
    centavos = int(round((numero - int(numero)) * 100))
    
    if numero < 1 and centavos > 0:
        return f"CERO PESOS CON {centavos}/100 M/CTE"
    
    if numero == 0:
        return 'CERO PESOS M/CTE'
    
    numero_str = str(numero).zfill(12)
    billones = convert_grupo(numero_str[:3])
    millones = convert_grupo(numero_str[3:6])
    miles = convert_grupo(numero_str[6:9])
    cientos = convert_grupo(numero_str[9:])
    
    texto = ''
    if billones:
        texto += billones + 'MIL '
    if millones:
        texto += millones + 'MILLONES '
    if miles:
        texto += miles + 'MIL '
    if cientos:
        texto += cientos
    
    texto += 'PESOS'
    
    if centavos > 0:
        texto += f" CON {centavos}/100"
    
    texto += ' M/CTE'
    
    return texto.strip()

def enviar_correo_orden_compra(orden, detalles, email_proveedor):
    try:
        msg = Message('Nueva Orden de Compra',
                      sender=current_app.config['MAIL_DEFAULT_SENDER'],
                      recipients=[email_proveedor])
        
        msg.body = f"""
        Se ha generado una nueva Orden de Compra.
        
        Número de Orden: {orden.Numero}
        Fecha: {orden.FechaCreacion.strftime('%Y-%m-%d %H:%M:%S')}
        
        Por favor, revise el archivo adjunto para más detalles.
        """
        
        pdf = generar_pdf_orden_compra(orden, detalles)
        msg.attach("orden_de_compra.pdf", "application/pdf", pdf)
        
        mail.send(msg)
        print(f"Correo enviado exitosamente a {email_proveedor}")
    except Exception as e:
        print(f"Error al enviar correo: {str(e)}")
        print(traceback.format_exc())

@app.route('/api/ultimo_consecutivo_ordenes_compra', methods=['GET', 'POST'])
def ultimo_consecutivo_ordenes_compra():
    try:
        consecutivo = Consecutivos.query.filter_by(Formulario='OC').first()
        if not consecutivo:
            return jsonify({'success': False, 'message': 'No se encontró el consecutivo para Órdenes de Compra'}), 404

        if request.method == 'GET':
            ultimo_consecutivo = f"{consecutivo.Prefijo}{consecutivo.Actual.zfill(2)}"
            return jsonify({'success': True, 'ultimoConsecutivo': ultimo_consecutivo})
        elif request.method == 'POST':
            # Incrementar el consecutivo
            actual = int(consecutivo.Actual)
            consecutivo.Actual = f"{actual + 1:02d}"
            db.session.commit()
            nuevo_consecutivo = f"{consecutivo.Prefijo}{consecutivo.Actual}"
            return jsonify({'success': True, 'nuevoConsecutivo': nuevo_consecutivo})
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'message': str(e)}), 500
    
@app.route('/api/ciudades', methods=['POST'])
@app.route('/api/ciudades/<string:id>', methods=['PUT'])
def manejar_ciudad(id=None):
    try:
        data = request.json
        app.logger.info(f"Datos recibidos: {data}")

        # Procesar porcreteica
        porcreteica = data.get('porcreteica')
        if porcreteica is not None and porcreteica != '':
            try:
                porcreteica = float(porcreteica)
            except ValueError:
                return jsonify({'error': 'El valor de porcreteica debe ser numérico'}), 400
        else:
            porcreteica = None

        if request.method == 'POST':
            # Verificar si la ciudad ya existe
            ciudad_existente = Ciudades.query.get(data['IdCiudad'])
            if ciudad_existente:
                return jsonify({'error': 'La ciudad ya existe'}), 400
            
            nueva_ciudad = Ciudades(
                IdCiudad=data['IdCiudad'],
                Ciudad=data['Ciudad'],
                IdDepartamento=data['IdDepartamento'],
                idpais=data['idpais'],
                porcreteica=porcreteica,
                Estado=data['Estado']
            )
            db.session.add(nueva_ciudad)
            mensaje = 'Ciudad creada exitosamente'
            codigo_respuesta = 201
        else:  # PUT
            ciudad = Ciudades.query.get(id)
            if not ciudad:
                return jsonify({'error': 'Ciudad no encontrada'}), 404
            ciudad.Ciudad = data['Ciudad']
            ciudad.IdDepartamento = data['IdDepartamento']
            ciudad.idpais = data['idpais']
            ciudad.porcreteica = porcreteica
            ciudad.Estado = data['Estado']
            mensaje = 'Ciudad actualizada exitosamente'
            codigo_respuesta = 200
        
        db.session.commit()
        app.logger.info(f"Operación exitosa: {mensaje}")
        return jsonify({'message': mensaje, 'ciudad': data}), codigo_respuesta
    except SQLAlchemyError as e:
        db.session.rollback()
        app.logger.error(f"Error de base de datos: {str(e)}")
        app.logger.error(traceback.format_exc())
        return jsonify({'error': 'Error de base de datos'}), 500
    except Exception as e:
        db.session.rollback()
        app.logger.error(f"Error al manejar ciudad: {str(e)}")
        app.logger.error(traceback.format_exc())
        return jsonify({'error': 'Error interno del servidor'}), 500

@app.route('/api/generar_imagen_licencia/<nit>', methods=['GET'])
def generar_imagen_licencia(nit):
    licencia = Licencia.query.filter_by(nit=nit).first()
    if licencia:
        try:
            img = Image.new('RGB', (800, 600), color = (255, 255, 255))
            d = ImageDraw.Draw(img)
            
            logo_path = os.path.join(app.root_path, 'static', 'img', 'Logo.png')
            logo = Image.open(logo_path)
            logo = logo.resize((100, 100))
            img.paste(logo, (10, 10), logo if logo.mode == 'RGBA' else None)
            
            firma_path = os.path.join(app.root_path, 'static', 'img', 'firma.png')
            firma = Image.open(firma_path)
            firma = firma.resize((150, 50))
            img.paste(firma, (600, 500), firma if firma.mode == 'RGBA' else None)

            font_path = os.path.join(app.root_path, 'static', 'fonts', 'arial.ttf')
            font = ImageFont.truetype(font_path, 16)

            d.text((10,120), f"www.migsistemas.com - info@migsistemas.com", font=font, fill=(0,0,0))
            d.text((10,140), f"Cel: 300 225 7898", font=font, fill=(0,0,0))
            d.text((10,180), f"Medellín, {datetime.now().strftime('%d de %B de %Y')}", font=font, fill=(0,0,0))
            d.text((10,220), f"Señores", font=font, fill=(0,0,0))
            d.text((10,240), f"{licencia.nombre_comercial}", font=font, fill=(0,0,0))
            d.text((10,260), f"{licencia.ciudad}", font=font, fill=(0,0,0))
            d.text((10,300), f"REF: Licenciamiento De Software MIG", font=font, fill=(0,0,0))
            d.text((10,340), f"NIT: {licencia.nit}", font=font, fill=(0,0,0))
            d.text((10,360), f"RAZÓN SOCIAL: {licencia.razon_social}", font=font, fill=(0,0,0))
            d.text((10,380), f"NOMBRE COMERCIAL: {licencia.nombre_comercial}", font=font, fill=(0,0,0))
            d.text((10,400), f"UBICACIÓN COMERCIAL: {licencia.ubicacion_comercial}", font=font, fill=(0,0,0))
            d.text((10,420), f"CIUDAD: {licencia.ciudad}", font=font, fill=(0,0,0))
            d.text((10,440), f"TELÉFONO: {licencia.telefono}", font=font, fill=(0,0,0))
            d.text((10,460), f"VERSIÓN: {licencia.version}", font=font, fill=(0,0,0))
            d.text((10,480), f"NUMERO DE LICENCIA: {licencia.numero_licencia}", font=font, fill=(0,0,0))
            d.text((10,500), f"CANTIDAD DE USUARIOS: {licencia.cantidad_usuarios}", font=font, fill=(0,0,0))
            d.text((10,520), f"TIPO LICENCIA: {licencia.tipo_licencia}", font=font, fill=(0,0,0))
            if licencia.fecha_vencimiento:
                d.text((10,540), f"FECHA DE VENCIMIENTO: {licencia.fecha_vencimiento}", font=font, fill=(0,0,0))

            d.text((10,580), "Lo anterior se expide para cumplimiento de requisitos exigidos por la ley,", font=font, fill=(0,0,0))
            d.text((10,600), "para legalidad y autenticidad.", font=font, fill=(0,0,0))
            d.text((10,640), "Cualquier duda con gusto la aclararemos.", font=font, fill=(0,0,0))

            d.text((600,560), "MARY PARRA G.", font=font, fill=(0,0,0))
            d.text((600,580), "SISTEMAS MIG S.A.S", font=font, fill=(0,0,0))
            d.text((600,600), "NIT: 900.275.400-8", font=font, fill=(0,0,0))

            img_io = io.BytesIO()
            img.save(img_io, 'PNG')
            img_io.seek(0)

            return send_file(img_io, mimetype='image/png')
        except Exception as e:
            logging.error(f"Error al generar la imagen de la licencia: {str(e)}")
            return jsonify({'success': False, 'message': 'Error al generar la imagen de la licencia'}), 500
    else:
        return jsonify({'success': False, 'message': 'No se encontró licencia para el NIT proporcionado'}), 404
    
@app.route('/test_email')
def test_email():
    try:
        msg = Message('Test Email desde Flask',
                    sender=app.config['MAIL_DEFAULT_SENDER'],
                    recipients=['riosjuan3053@gmail.com'])
        msg.body = "Este es un correo de prueba enviado desde la aplicación Flask."
        mail.send(msg)
        return "Correo de prueba enviado. Por favor, verifica tu bandeja de entrada y carpeta de spam."
    except Exception as e:
        return f"Error al enviar el correo de prueba: {str(e)}"

@app.errorhandler(404)
def not_found_error(error):
    return jsonify({'error': 'Recurso no encontrado'}), 404

@app.errorhandler(500)
def internal_error(error):
    db.session.rollback()
    return jsonify({'error': 'Error interno del servidor'}), 500

@app.route('/api/test', methods=['GET', 'POST'])
def test_route():
    if request.method == 'POST':
        # Manejar solicitud POST
        data = request.json
        return jsonify({"message": "POST recibido", "data": data})
    else:
        # Manejar solicitud GET
        return jsonify({"message": "Conexión exitosa"})

@app.route('/ping', methods=['GET'])
def ping():
    return "pong"

@app.route('/api/consecutivos', methods=['GET', 'POST', 'PUT'])
def manejar_consecutivos():
    def get_value(data, key, default=None):
        value = data.get(key)
        return value if value not in [None, "", "null", "undefined"] else default

    def obtener_iniciales(formulario):
        mapeo_formularios = {
            'Compras de Mercancía': 'CM',
            'Cotizaciones': 'COT',
            'Cuentas de Cobro': 'CC',
            'Devolución de Compras': 'DC',
            'Entradas de Inventario': 'EI',
            'Gastos': 'GAS',
            'Ordenes de Compra': 'OC',
            'Pedidos': 'PED',
            'Remisiones': 'REM',
            'Salidas': 'SAL',
            'Solicitud de Materiales': 'SM',
            'Traslados de Bodega': 'TB',
            'Inventario Físico': 'IF'
        }
        return mapeo_formularios.get(formulario, formulario[:3].upper())

    def obtener_nombre_completo(iniciales):
        mapeo_inverso = {
            'CM': 'Compras de Mercancía',
            'COT': 'Cotizaciones',
            'CC': 'Cuentas de Cobro',
            'DC': 'Devolución de Compras',
            'EI': 'Entradas de Inventario',
            'GAS': 'Gastos',
            'OC': 'Ordenes de Compra',
            'PED': 'Pedidos',
            'REM': 'Remisiones',
            'SAL': 'Salidas',
            'SM': 'Solicitud de Materiales',
            'TB': 'Traslados de Bodega',
            'IF': 'Inventario Físico'
        }
        return mapeo_inverso.get(iniciales, iniciales)

    if request.method in ['POST', 'PUT']:
        data = request.json
        print("Datos recibidos:", data)  # Log para depuración

        consecutivo_data = {
            'IdConsecutivo': int(get_value(data, 'IdConsecutivo')),
            'Consecutivo': get_value(data, 'Consecutivo'),  # Este es el campo 'Descripción'
            'Formulario': obtener_iniciales(get_value(data, 'Formulario')),
            'Prefijo': get_value(data, 'Prefijo'),
            'Desde': get_value(data, 'Desde'),
            'Hasta': get_value(data, 'Hasta'),
            'Actual': get_value(data, 'Actual'),
            'Resolucion': get_value(data, 'Resolucion'),
            'FechaResolucion': datetime.strptime(get_value(data, 'FechaInicioResolucion'), '%Y-%m-%d').date() if get_value(data, 'FechaInicioResolucion') else None,
            'ObservacionesResolucion': get_value(data, 'Observaciones'),
            'Estado': get_value(data, 'Activo', True),
            'Comprobante': get_value(data, 'TipoDocumentoFactura'),
            'fechafinresolucion': get_value(data, 'FechaFinResolucion'),
            'tiporesolucion': get_value(data, 'Tipo')
        }

        print("Datos procesados:", consecutivo_data)  # Log para depuración

        if request.method == 'POST':
            nuevo_consecutivo = Consecutivos(**consecutivo_data)
            db.session.add(nuevo_consecutivo)
        else:  # PUT
            consecutivo = Consecutivos.query.get(consecutivo_data['IdConsecutivo'])
            if consecutivo:
                for key, value in consecutivo_data.items():
                    setattr(consecutivo, key, value)
            else:
                return jsonify({'message': 'Consecutivo no encontrado'}), 404

        try:
            db.session.commit()
            return jsonify({'message': f'Consecutivo {"creado" if request.method == "POST" else "actualizado"} exitosamente'}), 200
        except Exception as e:
            db.session.rollback()
            return jsonify({'message': f'Error al {"crear" if request.method == "POST" else "actualizar"} el consecutivo: {str(e)}'}), 500

    if request.method == 'GET':
        consecutivos = Consecutivos.query.all()
        return jsonify([{
            'IdConsecutivo': c.IdConsecutivo,
            'Consecutivo': c.Consecutivo,
            'Formulario': obtener_nombre_completo(c.Formulario),
            'Prefijo': c.Prefijo,
            'Desde': c.Desde,
            'Hasta': c.Hasta,
            'Actual': c.Actual,
            'Resolucion': c.Resolucion,
            'FechaResolucion': c.FechaResolucion.strftime('%Y-%m-%d') if c.FechaResolucion else None,
            'ObservacionesResolucion': c.ObservacionesResolucion,
            'Estado': c.Estado,
            'Comprobante': c.Comprobante,
            'fechafinresolucion': c.fechafinresolucion,
            'tiporesolucion': c.tiporesolucion
        } for c in consecutivos])

@app.route('/api/consecutivos/<int:id>', methods=['GET', 'DELETE'])
def manejar_consecutivo_individual(id):
    consecutivo = Consecutivos.query.get_or_404(id)
    if request.method == 'GET':
        return jsonify({
            'IdConsecutivo': consecutivo.IdConsecutivo,
            'Consecutivo': consecutivo.Consecutivo,
            'Descripcion': consecutivo.Descripcion,
            'Formulario': consecutivo.Formulario,
            'Prefijo': consecutivo.Prefijo,
            'Desde': consecutivo.Desde,
            'Hasta': consecutivo.Hasta,
            'Actual': consecutivo.Actual,
            'Resolucion': consecutivo.Resolucion,
            'FechaInicioResolucion': consecutivo.FechaInicioResolucion.isoformat() if consecutivo.FechaInicioResolucion else None,
            'FechaFinResolucion': consecutivo.FechaFinResolucion.isoformat() if consecutivo.FechaFinResolucion else None,
            'TipoDocumentoFactura': consecutivo.TipoDocumentoFactura,
            'Observaciones': consecutivo.Observaciones,
            'Tipo': consecutivo.Tipo,
            'Estado': consecutivo.Estado
        })
    elif request.method == 'DELETE':
        try:
            db.session.delete(consecutivo)
            db.session.commit()
            return jsonify({'message': 'Consecutivo eliminado exitosamente'})
        except Exception as e:
            db.session.rollback()
            return jsonify({'message': f'Error al eliminar el consecutivo: {str(e)}'}), 500
        
@app.route('/api/consecutivo_entradas_inventario', methods=['GET'])
def obtener_consecutivo_entradas_inventario():
    consecutivo = Consecutivos.query.filter_by(Formulario='EI').first()
    if consecutivo:
        return jsonify({
            'IdConsecutivo': consecutivo.IdConsecutivo,
            'Consecutivo': consecutivo.Consecutivo,
            'Prefijo': consecutivo.Prefijo,
            'Actual': consecutivo.Actual
        })
    else:
        return jsonify({'error': 'No se encontró el consecutivo para Entradas de Inventario'}), 404

@app.route('/api/actualizar_consecutivo_entradas', methods=['POST'])
def actualizar_consecutivo_entradas():
    try:
        consecutivo = Consecutivos.query.filter_by(Formulario='EI').first()
        if consecutivo:
            actual = int(consecutivo.Actual)
            consecutivo.Actual = str(actual + 1).zfill(2)  # Usamos zfill(2) para tener siempre 2 dígitos
            db.session.commit()
            nuevo_consecutivo = f"{consecutivo.Prefijo}{consecutivo.Actual}"
            return jsonify({'success': True, 'nuevoConsecutivo': nuevo_consecutivo})
        else:
            return jsonify({'success': False, 'message': 'No se encontró el consecutivo para Entradas de Inventario'})
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'message': str(e)})
    
@app.route('/api/proveedores', methods=['GET'])
def get_proveedores():
    proveedores = Proveedores.query.all()
    return jsonify([{
        'Nit': p.Nit,
        'RazonSocial': p.RazonSocial,
        'Direccion': p.Direccion,
        'Telefono1': p.Telefono1,
        'Estado': p.Estado
    } for p in proveedores])
    
@app.route('/api/verificar_inventario', methods=['GET'])
def verificar_inventario():
    id_referencia = request.args.get('idReferencia')
    id_bodega = request.args.get('idBodega')
    mes_actual = datetime.now().strftime('%Y%m')

    saldo = SaldosBodega.query.filter_by(
        IdBodega=id_bodega,
        Mes=mes_actual,
        IdReferencia=id_referencia
    ).first()

    if saldo:
        saldo_disponible = saldo.SaldoInicial + saldo.Entradas + saldo.Compras - saldo.Salidas - saldo.Ventas
    else:
        saldo_disponible = 0

    return jsonify({
        'disponible': saldo_disponible > 0,
        'saldoDisponible': float(saldo_disponible)
    })
    
@app.route('/api/referencias', methods=['POST', 'PUT'])
def manejar_referencias():
    data = request.json
    
    # Calcular el precio con IVA
    precio_sin_iva = data.get('precioVenta1')
    iva_porcentaje = data.get('iva', 0)
    if precio_sin_iva is not None and iva_porcentaje is not None:
        precio_sin_iva = float(precio_sin_iva)
        iva_porcentaje = float(iva_porcentaje)
        precio_con_iva = precio_sin_iva * (1 + (iva_porcentaje / 100))
    else:
        precio_con_iva = None
    
    referencia_data = {
        'IdReferencia': data.get('codigo'),
        'Referencia': data.get('descripcion'),
        'IdGrupo': data.get('grupo'),
        'IdUnidad': data.get('unidad'),
        'Ubicacion': data.get('ubicacion') or None,
        'Costo': data.get('costo') or None,
        'PrecioVenta1': precio_con_iva,
        'IVA': iva_porcentaje,
        'Marca': data.get('marca') or None,
        'EstadoProducto': data.get('estadoProducto') or None,
        'Estado': data.get('activo', True),
        'Tipo': data.get('esServicio', False),
        'ManejaInventario': not data.get('esServicio', False),
        'productoagotado': data.get('agotado', False),
        'modificaprecio': data.get('modificaPrecio', False),
        'idsubgrupo': data.get('subgrupo') or None,
        'idsubcategoria': data.get('subcategoria') or None,
        'idbodega': data.get('bodega') or None
    }

    try:
        referencia = Referencia.query.get(referencia_data['IdReferencia'])
        if referencia:
            # Si la referencia existe, actualizar sus campos
            for key, value in referencia_data.items():
                setattr(referencia, key, value)
            mensaje = 'Referencia actualizada exitosamente'
        else:
            # Si la referencia no existe, crear una nueva
            nueva_referencia = Referencia(**referencia_data)
            db.session.add(nueva_referencia)
            mensaje = 'Referencia creada exitosamente'

        db.session.commit()
        return jsonify({'message': mensaje}), 200
    except SQLAlchemyError as e:
        db.session.rollback()
        return jsonify({'message': f'Error en la base de datos: {str(e)}'}), 500
    except Exception as e:
        db.session.rollback()
        return jsonify({'message': f'Error inesperado: {str(e)}'}), 500

@app.route('/api/referencias', methods=['GET'])
def obtener_referencias():
    filtro = request.args.get('filtro', '')
    id_bodega = request.args.get('idBodega', '')
    mes_actual = datetime.now().strftime('%Y%m')

    query = text("""
        SELECT 
            r."IdReferencia", 
            r."Referencia", 
            r."PrecioVenta1", 
            r."IVA", 
            r."Ubicacion", 
            r."idbodega", 
            r."IdUnidad", 
            COALESCE(s."Saldo", 0) as Saldo
        FROM 
            referencias r
        LEFT JOIN 
            "SaldosBodega" s
        ON 
            r."IdReferencia" = s."IdReferencia"
            AND s."IdBodega" = :id_bodega
            AND s."Mes" = :mes_actual
        WHERE 
            (r."IdReferencia" ILIKE :filtro OR r."Referencia" ILIKE :filtro)
            AND r."Estado" = True
    """)

    try:
        result = db.session.execute(query, {
            'filtro': f'%{filtro}%',
            'id_bodega': id_bodega,
            'mes_actual': mes_actual
        })
        
        referencias = [{
            'IdReferencia': row.IdReferencia,
            'Referencia': row.Referencia,
            'PrecioVenta1': str(row.PrecioVenta1),
            'IVA': str(row.IVA),
            'Ubicacion': row.Ubicacion,
            'idbodega': row.idbodega,
            'IdUnidad': row.IdUnidad,
            'Saldo': str(row.Saldo)
        } for row in result]

        return jsonify(referencias)
    except Exception as e:
        print(f"Error al obtener referencias: {str(e)}")
        return jsonify({'error': 'Error al obtener referencias'}), 500

@app.route('/api/salidas_inventario', methods=['POST'])
def crear_salida_inventario():
    data = request.json
    
    nueva_salida = Salidas1(
        Numero=data['Numero'],
        Mes=data['Mes'],
        IdBodega=data['IdBodega'],
        Observaciones=data['Observaciones'],
        FechaCreacion=datetime.now(),
        IdUsuario=data['IdUsuario'],
        IdConsecutivo=data['IdConsecutivo'],
        Total=data['Total']
    )
    db.session.add(nueva_salida)
    
    for detalle in data['detalles']:
        nuevo_detalle = Salidas2(
            ID=f"{data['Numero']}_{detalle['IdReferencia']}",
            Numero=data['Numero'],
            IdReferencia=detalle['IdReferencia'],
            Descripcion=detalle['Descripcion'],
            Cantidad=detalle['Cantidad'],
            Valor=detalle['Valor'],
            Subtotal=detalle['Subtotal'],
            Idunidad=detalle['Idunidad']
        )
        db.session.add(nuevo_detalle)
    
    db.session.commit()
    return jsonify({'success': True, 'message': 'Salida de inventario creada exitosamente'}), 201

@app.route('/api/guardar_salida', methods=['POST'])
def guardar_salida():
    data = request.json
    salida1_data = data.get('salida1', {})
    salidas2_data = data.get('salidas2', [])

    try:
        # Asegurarse de que el usuario 'MIG' existe
        get_or_create_mig_user()

        # Crear instancia de Salidas1
        nueva_salida1 = Salidas1(
            Numero=salida1_data['Numero'],
            Mes=salida1_data['Mes'],
            Anulado=salida1_data['Anulado'],
            IdBodega=salida1_data['IdBodega'],
            CuentaDebito=salida1_data.get('CuentaDebito'),
            CuentaCredito=salida1_data.get('CuentaCredito'),
            Observaciones=salida1_data.get('Observaciones'),
            FechaCreacion=datetime.fromisoformat(salida1_data['FechaCreacion'].replace('Z', '+00:00')),
            IdUsuario='MIG',
            Recibe=salida1_data.get('Recibe'),
            idproyecto=salida1_data.get('idproyecto'),
            fechamodificacion=datetime.fromisoformat(salida1_data['fechamodificacion'].replace('Z', '+00:00')) if salida1_data.get('fechamodificacion') else None,
            IdConsecutivo=salida1_data['IdConsecutivo'],
            op=datetime.fromisoformat(salida1_data.get('op').replace('Z', '+00:00')).strftime('%Y-%m-%d'),
            fecha=datetime.fromisoformat(salida1_data['fecha'].replace('Z', '+00:00')),
            subtotal=Decimal(str(salida1_data.get('subtotal', '0'))),
            total_iva=Decimal(str(salida1_data.get('total_iva', '0'))),
            total_impoconsumo=Decimal(str(salida1_data.get('total_impoconsumo', '0'))),
            total_ipc=Decimal(str(salida1_data.get('total_ipc', '0'))),
            total_ibua=Decimal(str(salida1_data.get('total_ibua', '0'))),
            total_icui=Decimal(str(salida1_data.get('total_icui', '0'))),
            total=Decimal(str(salida1_data.get('total', '0')))
        )
        db.session.add(nueva_salida1)

        # Guardar en Salidas2 y actualizar SaldosBodega
        mes_actual = datetime.now().strftime('%Y%m')
        for salida2 in salidas2_data:
            if not salida2['IdReferencia']:
                logging.warning(f"Referencia vacía encontrada en salidas2_data: {salida2}")
                continue  # Saltar referencias vacías

            # Verificar si la referencia existe
            referencia = Referencia.query.get(salida2['IdReferencia'])
            if not referencia:
                logging.error(f"Referencia no encontrada: {salida2['IdReferencia']}")
                return jsonify({'success': False, 'message': f'La referencia {salida2["IdReferencia"]} no existe'}), 400

            nueva_salida2 = Salidas2(
                ID=salida2['ID'],
                Numero=salida2['Numero'],
                IdReferencia=salida2['IdReferencia'],
                Descripcion=salida2['Descripcion'],
                Cantidad=Decimal(str(salida2['Cantidad'])),
                Valor=Decimal(str(salida2['Valor'])),
                IVA=Decimal(str(salida2.get('IVA', '0'))),
                Descuento=Decimal('0'),
                lote=salida2.get('lote', ''),
                idunidad=salida2['idunidad'],
                impoconsumo=Decimal(str(salida2.get('impoconsumo', '0'))),
                ipc=Decimal(str(salida2.get('ipc', '0'))),
                imp_ibua=Decimal(str(salida2.get('imp_ibua', '0'))),
                imp_icui=Decimal(str(salida2.get('imp_icui', '0')))
            )
            db.session.add(nueva_salida2)

            # Actualizar SaldosBodega
            saldo = SaldosBodega.query.filter_by(
                IdBodega=salida1_data['IdBodega'],
                Mes=mes_actual,
                IdReferencia=salida2['IdReferencia']
            ).first()

            if saldo:
                saldo.Salidas += Decimal(str(salida2['Cantidad']))
                saldo.Saldo -= Decimal(str(salida2['Cantidad']))
            else:
                nuevo_saldo = SaldosBodega(
                    IdBodega=salida1_data['IdBodega'],
                    Mes=mes_actual,
                    IdReferencia=salida2['IdReferencia'],
                    Salidas=Decimal(str(salida2['Cantidad'])),
                    Saldo=-Decimal(str(salida2['Cantidad'])),
                    SaldoInicial=Decimal('0'),
                    Entradas=Decimal('0'),
                    Compras=Decimal('0'),
                    Ventas=Decimal('0')
                )
                db.session.add(nuevo_saldo)

        db.session.commit()

        # Generar documento Word en memoria
        doc_content = generar_documento_salida(salida1_data, salidas2_data)
        filename = f"Salida_Bodega_{salida1_data['Numero']}.docx"

        return jsonify({
            'success': True, 
            'message': 'Salida guardada y saldo actualizado con éxito',
            'documento': filename,
            'doc_content': base64.b64encode(doc_content).decode('utf-8')
        })
    except SQLAlchemyError as e:
        db.session.rollback()
        error_msg = f"Error de base de datos al guardar la salida: {str(e)}"
        logging.error(error_msg)
        logging.error(traceback.format_exc())
        return jsonify({'success': False, 'message': error_msg}), 500
    except Exception as e:
        db.session.rollback()
        error_msg = f"Error inesperado al guardar la salida: {str(e)}"
        logging.error(error_msg)
        logging.error(traceback.format_exc())
        return jsonify({'success': False, 'message': error_msg}), 500

@app.route('/api/verificar_referencias', methods=['POST'])
def verificar_referencias():
    try:
        data = request.json
        referencias = data.get('referencias', [])
        
        app.logger.info(f"Referencias recibidas para verificar: {referencias}")
        
        if not referencias:
            app.logger.warning("No se proporcionaron referencias para verificar")
            return jsonify({
                'success': False,
                'message': 'No se proporcionaron referencias para verificar'
            }), 400

        invalid_references = []
        for referencia in referencias:
            # Limpia la referencia de caracteres no deseados
            referencia_limpia = ''.join(char for char in referencia if char.isalnum() or char in ['-', '_'])
            app.logger.debug(f"Verificando referencia: {referencia_limpia}")
            if not Referencia.query.get(referencia_limpia):
                invalid_references.append(referencia_limpia)
        
        if invalid_references:
            app.logger.warning(f"Referencias inválidas encontradas: {invalid_references}")
            return jsonify({
                'success': False,
                'message': 'Algunas referencias no existen',
                'invalid_references': invalid_references
            }), 400
        
        app.logger.info("Todas las referencias son válidas")
        return jsonify({
            'success': True,
            'message': 'Todas las referencias son válidas'
        })

    except SQLAlchemyError as e:
        db.session.rollback()
        app.logger.error(f"Error de base de datos al verificar referencias: {str(e)}")
        return jsonify({
            'success': False,
            'message': 'Error de base de datos al verificar referencias',
            'error': str(e)
        }), 500
    except Exception as e:
        app.logger.error(f"Error inesperado al verificar referencias: {str(e)}")
        return jsonify({
            'success': False,
            'message': 'Error inesperado al verificar referencias',
            'error': str(e)
        }), 500

def generar_documento_salida(salida, productos):
    doc = Document()

    # Configurar márgenes de página
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

    # Crear tabla para el encabezado
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.allow_autofit = False
    
    # Ajustar ancho de las columnas
    table.columns[0].width = Inches(2.5)
    table.columns[1].width = Inches(4.5)

    # Celda para el logo
    logo_cell = table.cell(0, 0)
    logo_path = os.path.join(current_app.root_path, 'static', 'img', 'logoEmpresa.png')
    if os.path.exists(logo_path):
        logo_paragraph = logo_cell.paragraphs[0]
        logo_run = logo_paragraph.add_run()
        logo_run.add_picture(logo_path, width=Inches(2))
    else:
        print(f"Logo no encontrado en: {logo_path}")
    
    company_info = logo_cell.add_paragraph()
    company_info.add_run('CCD INGENIERÍA Y CONSTRUCCIONES S.A.S.\n').bold = True
    company_info.add_run('NIT. 901.092.189-5')

    # Celda para el título y la fecha
    title_cell = table.cell(0, 1)
    title_paragraph = title_cell.paragraphs[0]
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    title_run = title_paragraph.add_run('Salida de Bodega\n')
    title_run.bold = True
    title_run.font.size = Pt(16)
    date_run = title_paragraph.add_run(f"Fecha: {salida.get('fecha', datetime.now().strftime('%Y-%m-%d'))}")
    date_run.font.size = Pt(10)

    # Agregar línea separadora
    doc.add_paragraph('_' * 100)

    # Resto del documento
    doc.add_heading('Observaciones', level=1)
    doc.add_paragraph(salida.get('Observaciones', 'Sin observaciones'))

    doc.add_heading('Productos', level=1)
    products_table = doc.add_table(rows=1, cols=5)
    products_table.style = 'Table Grid'
    hdr_cells = products_table.rows[0].cells
    hdr_cells[0].text = 'ID Referencia'
    hdr_cells[1].text = 'Descripción'
    hdr_cells[2].text = 'Cantidad'
    hdr_cells[3].text = 'Valor'
    hdr_cells[4].text = 'Subtotal'

    for producto in productos:
        row_cells = products_table.add_row().cells
        row_cells[0].text = producto.get('IdReferencia', '')
        row_cells[1].text = producto.get('Descripcion', '')
        row_cells[2].text = str(producto.get('Cantidad', 0))
        row_cells[3].text = str(producto.get('Valor', 0))
        subtotal = float(producto.get('Cantidad', 0)) * float(producto.get('Valor', 0))
        row_cells[4].text = str(subtotal)

    doc.add_heading('Firma', level=1)
    doc.add_paragraph('_' * 30)
    doc.add_paragraph('Nombre y Firma')

    # Guardar en un objeto BytesIO
    from io import BytesIO
    doc_buffer = BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)
    return doc_buffer.getvalue()

@app.route('/api/salidas_inventario/<numero>', methods=['PUT'])
def actualizar_salida_inventario(numero):
    data = request.json
    
    salida = Salidas1.query.filter_by(Numero=numero).first()
    if not salida:
        return jsonify({'success': False, 'message': 'Salida no encontrada'}), 404
    
    salida.Total = data['Total']
    salida.Observaciones = data['Observaciones']
    
    Salidas2.query.filter_by(Numero=numero).delete()
    
    for detalle in data['detalles']:
        nuevo_detalle = Salidas2(
            ID=f"{numero}_{detalle['IdReferencia']}",
            Numero=numero,
            IdReferencia=detalle['IdReferencia'],
            Descripcion=detalle['Descripcion'],
            Cantidad=detalle['Cantidad'],
            Valor=detalle['Valor'],
            Subtotal=detalle['Subtotal'],
            Idunidad=detalle['Idunidad']
        )
        db.session.add(nuevo_detalle)
    
    db.session.commit()
    return jsonify({'success': True, 'message': 'Salida de inventario actualizada exitosamente'})

# Función auxiliar para crear o obtener el usuario 'MIG'
def get_or_create_mig_user():
    mig_user = Usuarios.query.get('MIG')
    if not mig_user:
        mig_user = Usuarios(
            IdUsuario='MIG',
            Contraseña='password_temporal',  # Asegúrate de cambiar esto después
            Descripcion='Usuario por defecto',
            Estado=True
        )
        db.session.add(mig_user)
        db.session.commit()
    return mig_user

@app.route('/api/entradas_informes', methods=['GET'])
@cross_origin()
def get_entradas_informes():
    app.logger.info('Función get_entradas_informes llamada')
    fecha_inicio = request.args.get('fecha_inicio')
    fecha_fin = request.args.get('fecha_fin')
    
    app.logger.info(f'Fechas recibidas: inicio={fecha_inicio}, fin={fecha_fin}')

    if not fecha_inicio or not fecha_fin:
        return jsonify({'error': 'Se requieren fechas de inicio y fin'}), 400

    try:
        fecha_inicio = datetime.strptime(fecha_inicio, '%Y-%m-%d').date()
        fecha_fin = datetime.strptime(fecha_fin, '%Y-%m-%d').date()
    except ValueError:
        return jsonify({'error': 'Formato de fecha inválido'}), 400

    try:
        # Primero, obtenemos los datos de Entradas1
        entradas1 = Entradas1.query.filter(Entradas1.fecha.between(fecha_inicio, fecha_fin)).all()
        
        resultado = []
        for e1 in entradas1:
            # Luego, para cada entrada en Entradas1, buscamos las correspondientes en Entradas2
            entradas2 = Entradas2.query.filter_by(Numero=e1.Numero).all()
            for e2 in entradas2:
                resultado.append({
                    'Numero': e1.Numero,
                    'Fecha': e1.FechaCreacion.strftime('%Y-%m-%d %H:%M:%S') if e1.FechaCreacion else None,
                    'IdReferencia': e2.IdReferencia,
                    'Descripcion': e2.Descripcion,
                    'Cantidad': float(e2.Cantidad) if e2.Cantidad is not None else None,
                    'Valor': float(e2.Valor) if e2.Valor is not None else None,
                    'Total': float(e1.total) if e1.total is not None else None
                })

        app.logger.info(f'Número de entradas encontradas: {len(resultado)}')
        
        return jsonify({
            'entradas': resultado,
            'total_entradas': len(resultado)
        })
    except Exception as e:
        app.logger.error(f'Error en get_entradas_informes: {str(e)}')
        return jsonify({'error': f'Error al obtener los datos: {str(e)}'}), 500

@app.route('/api/actualizar_consecutivo_salidas_inventario', methods=['POST'])
def actualizar_consecutivo_salidas_inventario():
    try:
        consecutivo = Consecutivos.query.filter_by(Formulario='SAL').first()
        if consecutivo:
            # Incrementar el consecutivo
            actual = int(consecutivo.Actual)
            actual += 1
            # Formatear el nuevo consecutivo con solo dos dígitos
            consecutivo.Actual = f"{actual:02d}"
            db.session.commit()
            
            nuevo_consecutivo = f"{consecutivo.Prefijo}{consecutivo.Actual}"
            return jsonify({'success': True, 'nuevoConsecutivo': nuevo_consecutivo})
        else:
            return jsonify({'success': False, 'message': 'No se encontró el consecutivo para Salidas de Inventario'}), 404
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'message': str(e)}), 500
    
@app.route('/api/referencias', methods=['POST'])
def crear_referencia():
    data = request.json
    logging.debug(f"Datos recibidos: {data}")
    
    try:
        # Usar el código generado como IdReferencia
        nuevo_id = data.get('codigo')
        if not nuevo_id:
            return jsonify({'error': 'El código del producto es requerido'}), 400

        nueva_referencia = Referencia(
            IdReferencia=nuevo_id,
            Referencia=data['descripcion'],
            IdGrupo=data['grupo'],
            IdUnidad=data['unidad'],
            Costo=data.get('costo', 0),
            PrecioVenta1=data.get('precioVenta1', 0),
            IVA=data.get('iva', 0),
            Ubicacion=data.get('ubicacion'),
            Marca=data.get('marca'),
            EstadoProducto=data.get('estadoProducto', 'Bueno'),
            Estado=data.get('activo', True),
            Tipo=data.get('esServicio', False),
            ManejaInventario=not data.get('esServicio', False),
            productoagotado=data.get('agotado', False),
            idsubgrupo=data.get('subgrupo'),
            idsubcategoria=data.get('subcategoria'),
            idbodega=data.get('bodega'),
            FechaCreacion=datetime.utcnow(),
            StockMinimo=0,
            StockMaximo=0,
            SaldoAntesInv=0,
            Insumo=False,
            costoreal=0
        )
        
        logging.debug(f"Nueva referencia creada: {nueva_referencia.__dict__}")
        
        db.session.add(nueva_referencia)
        db.session.commit()
        
        logging.info(f"Referencia creada exitosamente con código: {nuevo_id}")
        return jsonify({
            'message': 'Referencia creada exitosamente',
            'codigo': nuevo_id
        }), 201

    except SQLAlchemyError as e:
        db.session.rollback()
        logging.error(f"Error de SQLAlchemy: {str(e)}")
        logging.error(f"Detalles del error: {e.__dict__}")
        return jsonify({'error': f'Error de base de datos: {str(e)}'}), 400
    except Exception as e:
        db.session.rollback()
        logging.error(f"Error inesperado: {str(e)}")
        logging.error(f"Detalles del error: {traceback.format_exc()}")
        return jsonify({'error': f'Error inesperado: {str(e)}'}), 500

@app.route('/api/grupos/<string:id_grupo>/siguiente-codigo', methods=['POST'])
def obtener_siguiente_codigo(id_grupo):
    try:
        grupo = Grupo.query.get(id_grupo)
        if not grupo:
            return jsonify({'error': 'Grupo no encontrado'}), 404

        # Incrementar el último código
        if grupo.ultimoCodigo is None:
            grupo.ultimoCodigo = 0
        grupo.ultimoCodigo += 1

        # Generar el nuevo código
        nuevo_codigo = f"{id_grupo}{grupo.ultimoCodigo:02d}"

        # Guardar los cambios
        db.session.commit()

        return jsonify({'nuevoCodigo': nuevo_codigo}), 200
    except Exception as e:
        db.session.rollback()
        logging.error(f"Error al generar siguiente código: {str(e)}")
        return jsonify({'error': 'Error al generar el código'}), 500

if __name__ == '__main__':
    app.run(debug=True)